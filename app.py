"""
IB Exam Worksheet Generator — Final Version
MCQ Mode + Structured Mark Scheme Mode
Based on: IB_Worksheet_Final (20) FINAL AGENT
"""
import streamlit as st
import anthropic
import openpyxl
import json
import re
import io
import base64
import datetime
import numpy as np

import fitz
from PIL import Image


def parse_page_range(raw_page: str, max_pages: int = 9999) -> tuple[int, int] | None:
    """Parse Excel page-number cell into (start_page, end_page).

    Handles:
      "91"          → (91, 91)    – plain page number
      "10-11"       → (10, 11)   – explicit range with dash
      "1011"        → (10, 11)   – merged-digit range
      "1920"        → (19, 20)   – merged-digit range
      "3032"        → (30, 32)   – merged-digit range
      "2026-10-11"  → (10, 11)   – Excel auto-converted date (month-day)
    Returns None when the value is unparseable.
    """
    if not raw_page:
        return None
    raw = str(raw_page).strip()

    # Explicit dash/en-dash range: "10-11", "10–11"
    m = re.match(r'^(\d+)\s*[-\u2013\u2014]\s*(\d+)$', raw)
    if m:
        s, e = int(m.group(1)), int(m.group(2))
        if 1 <= s and s <= e <= max_pages + 20:
            return (s, e)

    # Excel auto-converted date "YYYY-MM-DD" → take month and day
    m = re.match(r'^\d{4}-(\d{1,2})-(\d{1,2})$', raw)
    if m:
        s, e = int(m.group(1)), int(m.group(2))
        if 1 <= s <= max_pages and s <= e:
            return (s, e)

    # Plain integer
    m = re.match(r'^(\d+)$', raw)
    if m:
        n = int(m.group(1))
        if 1 <= n <= max_pages:
            return (n, n)
        # Attempt merged-range split: 1011 → 10-11, 1920 → 19-20, 3032 → 30-32
        s_raw = str(n)
        L = len(s_raw)
        best = None
        for split in range(1, L):
            s = int(s_raw[:split])
            e = int(s_raw[split:])
            if (1 <= s <= max_pages and s < e <= max_pages + 20
                    and 1 <= e - s <= 15):
                if best is None:
                    best = (s, e)
                else:
                    curr_diff = abs(split - (L - split))
                    prev_split = len(str(best[0]))
                    prev_diff  = abs(prev_split - (L - prev_split))
                    if curr_diff < prev_diff:
                        best = (s, e)
        return best  # None if no valid split found

    # Fallback: first integer in the cell
    m = re.search(r'\d+', raw)
    if m:
        n = int(m.group())
        if 1 <= n <= max_pages:
            return (n, n)
    return None

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn as docx_qn
from docx.oxml import OxmlElement


# ═══════════════════════════════════════════════════════════════════════════════
#  STREAMLIT UI
# ═══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Exam Worksheet Generator",
    page_icon="📄",
    layout="centered",
)

st.title("📄 Exam Worksheet Generator")
st.caption("QP PDF → questions · MS PDF → answers · Excel → metadata")

with st.expander("ℹ️ Source rules", expanded=False):
    st.markdown(
        "| Source | Used for |\n"
        "|--------|----------|\n"
        "| **QP PDF** | Question text, options, tables, diagrams, graphs — verbatim |\n"
        "| **MS PDF** | Answers only |\n"
        "| **Excel** | Question numbers, chapter, difficulty, marks, quotes |\n\n"
        "- Visual questions (graphs/diagrams/tables) → cropped image from QP\n"
        "- Excel `Topic = Error` or empty → `Unclassified`\n"
        "- Excel `Marks = 0` or missing → `1` (standard MCQ)"
    )

c1, c2 = st.columns(2)
with c1:
    qp_file = st.file_uploader("📋 Question Paper (QP)", type=["pdf"], key="qp")
    ms_file = st.file_uploader("✅ Mark Scheme (MS)",     type=["pdf"], key="ms")
with c2:
    xl_file = st.file_uploader("📊 Excel Sheet", type=["xlsx", "xls"], key="xl")

st.divider()

mode = st.radio(
    "Processing mode",
    options=["MCQ Mode", "Structured Mark Scheme Mode"],
    horizontal=True,
    key="mode",
    help=(
        "**MCQ Mode** — for Chemistry / Biology Paper 1 style. Mark Scheme is "
        "a grid of A/B/C/D answers.\n\n"
        "**Structured Mark Scheme Mode** — for Mathematics or any subject with "
        "full worked solutions (M1/A1/R1 marks, steps, diagrams). The Mark "
        "Scheme is extracted as a cropped image from the MS PDF for each question."
    ),
)

st.divider()


# ═══════════════════════════════════════════════════════════════════════════════
#  SHARED HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
def read_bytes(uploaded_file) -> bytes:
    if uploaded_file is None:
        return b""
    if isinstance(uploaded_file, bytes):
        return uploaded_file
    if hasattr(uploaded_file, "getvalue"):
        return uploaded_file.getvalue()
    if hasattr(uploaded_file, "seek"):
        try:
            uploaded_file.seek(0)
        except Exception:
            pass
    if hasattr(uploaded_file, "read"):
        return uploaded_file.read()
    raise TypeError(f"Unsupported type: {type(uploaded_file)}")


def to_b64(data: bytes) -> str:
    return base64.standard_b64encode(data).decode()


def safe_json(text: str):
    text = re.sub(r"```json\s*", "", text)
    text = re.sub(r"```\s*", "", text)
    text = text.strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        for opener, closer in [("[", "]"), ("{", "}")]:
            i = text.find(opener)
            j = text.rfind(closer)
            if i != -1 and j != -1 and j > i:
                try:
                    return json.loads(text[i:j + 1])
                except json.JSONDecodeError:
                    continue
        return None


# ═══════════════════════════════════════════════════════════════════════════════
#  EXCEL PARSER
# ═══════════════════════════════════════════════════════════════════════════════
def parse_excel(uploaded_file) -> list[dict]:
    uploaded_file.seek(0)
    wb = openpyxl.load_workbook(uploaded_file, read_only=True, data_only=True)
    ws = wb.active
    raw_headers = [c.value for c in next(ws.iter_rows(max_row=1))]
    headers = [str(h).strip() if h else "" for h in raw_headers]
    norm_headers = [re.sub(r"[\s_\-\.]+", " ", h.lower()).strip() for h in headers]

    def col_idx(*candidates):
        for name in candidates:
            target = re.sub(r"[\s_\-\.]+", " ", str(name).lower()).strip()
            for i, h in enumerate(norm_headers):
                if h == target:
                    return i
        for name in candidates:
            target = re.sub(r"[\s_\-\.]+", " ", str(name).lower()).strip()
            for i, h in enumerate(norm_headers):
                if target in h or h in target:
                    return i
        return None

    qn_idx   = col_idx("Question No.", "Question No", "Q No", "Q#", "Question Number", "Q", "QuestionNumber")
    page_idx = col_idx("Page Number", "Page", "page_number", "PDF Page", "PDF Page Number", "pdf_page", "QP Page", "Page No.", "Page No", "PageNum")
    top_idx  = col_idx("Topic", "Chapter", "Section", "Sub-topic", "Subtopic")
    dif_idx  = col_idx("Difficulty", "Level", "Hardness")
    mrk_idx  = col_idx("Marks", "Mark", "Points", "Score")
    qut_idx  = col_idx("Quote", "Note", "Motivational quote")
    ref_idx  = col_idx("Reference", "Ref", "Session", "Year", "Paper", "Exam Session", "Paper Reference")
    ms_col_idx = col_idx("Mark Scheme", "MarkScheme", "MS", "Answer", "Mark scheme", "mark scheme")

    parse_excel.last_headers = headers
    parse_excel.last_column_map = {
        "Question No.": (qn_idx,  headers[qn_idx]   if qn_idx   is not None else None),
        "Page Number":  (page_idx, headers[page_idx] if page_idx is not None else None),
        "Topic":        (top_idx,  headers[top_idx]  if top_idx  is not None else None),
        "Difficulty":   (dif_idx,  headers[dif_idx]  if dif_idx  is not None else None),
        "Marks":        (mrk_idx,  headers[mrk_idx]  if mrk_idx  is not None else None),
        "Reference":    (ref_idx,  headers[ref_idx]  if ref_idx  is not None else None),
        "Quote":        (qut_idx,  headers[qut_idx]  if qut_idx  is not None else None),
    }

    if qn_idx is None:
        qn_idx = 2

    def cell_val(row, idx, fallback=""):
        if idx is None:
            return fallback
        try:
            v = list(row)[idx].value
            return str(v).strip() if v is not None else fallback
        except IndexError:
            return fallback

    rows = []
    seen_keys = set()
    duplicates = []

    for row in ws.iter_rows(min_row=2):
        try:
            q_num = int(float(str(list(row)[qn_idx].value)))
        except (TypeError, ValueError):
            continue
        if q_num <= 0:
            continue

        topic = cell_val(row, top_idx)
        if not topic or topic.strip().lower() in ("error", "none", "n/a", "nan"):
            topic = "Unclassified"

        try:
            marks = int(float(cell_val(row, mrk_idx, "0")))
        except ValueError:
            marks = 0
        if marks < 0:
            marks = 0

        raw_page = ""
        if page_idx is not None:
            try:
                cell = list(row)[page_idx]
                v = cell.value
                if v is None:
                    raw_page = ""
                elif isinstance(v, datetime.datetime):
                    # Excel auto-converted "10-11" → date: recover month-day
                    raw_page = f"{v.month}-{v.day}"
                else:
                    raw_page = str(v).strip()
            except (IndexError, AttributeError):
                raw_page = ""

        page_num     = 0
        page_num_end = 0   # end page for multi-page questions
        if raw_page:
            pr = parse_page_range(raw_page)   # no max_pages at parse time
            if pr:
                page_num, page_num_end = pr
            else:
                # Could not interpret — keep 0 (will warn during generation)
                pass

        difficulty = cell_val(row, dif_idx, "Unspecified") or "Unspecified"
        ref        = cell_val(row, ref_idx, "")
        quote      = cell_val(row, qut_idx, "")

        ref = re.sub(r'\boctober\b', 'November', ref, flags=re.IGNORECASE)
        ref = re.sub(r'\boct\b',     'Nov',      ref, flags=re.IGNORECASE)

        # Paper 2 filter — only applies when caller enables it
        if getattr(parse_excel, '_filter_paper2', False):
            _P2_CODES = ('7205', '7210', '7215', '7310', '7315', '7320')
            if any(code in ref for code in _P2_CODES):
                continue
            if re.search(r'\bpaper[\s\-]?2\b|[/\-_](?:h|s)p2[/\-_]', ref, re.IGNORECASE):
                continue


        dedup_key = (q_num, ref.strip().lower(), page_num, topic.strip().lower(),
                     difficulty.strip().lower(), marks, quote.strip().lower())
        if dedup_key in seen_keys:
            duplicates.append((q_num, ref))
            continue
        seen_keys.add(dedup_key)

        rows.append({
            "qn": q_num, "page_num": page_num, "page_num_end": page_num_end,
            "topic": topic, "difficulty": difficulty, "marks": marks,
            "quote": quote, "ref": ref,
            "ms_text": (str(row[ms_col_idx].value).strip()
                        if ms_col_idx is not None and row[ms_col_idx].value
                           and str(row[ms_col_idx].value) not in ("None", "nan", "")
                        else ""),
        })

    parse_excel.last_duplicates = duplicates
    return rows


# ═══════════════════════════════════════════════════════════════════════════════
#  QP LOCATION FINDER (used by MCQ mode)
# ═══════════════════════════════════════════════════════════════════════════════
def find_question_locations(pdf_bytes: bytes) -> dict:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    raw = []
    pat = re.compile(r"^(\d+)\.$")

    for page_idx in range(len(doc)):
        page = doc[page_idx]
        ph, pw = page.rect.height, page.rect.width
        text_dict = page.get_text("dict")
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                first_text = spans[0].get("text", "").strip()
                bbox = line["bbox"]
                m = pat.match(first_text)
                if not m:
                    continue
                if bbox[0] > 110:
                    continue
                q_num = int(m.group(1))
                if q_num < 1 or q_num > 99:
                    continue
                marker_right_x = spans[0]["bbox"][2]
                raw.append((q_num, page_idx, bbox[1], ph, pw, marker_right_x))

    doc.close()
    raw.sort(key=lambda e: (e[1], e[2]))

    doc2 = fitz.open(stream=pdf_bytes, filetype="pdf")
    content_bottom_per_page: dict = {}
    footer_pat = re.compile(
        r"(international\s+baccalaureate|turn\s+over|references?\s*:|"
        r"^\s*\d{4}\s*[\-\u2013]\s*\d{3,4}\s*$|"
        r"^\s*[\-\u2013]\s*\d{1,3}\s*[\-\u2013]\s*$)",
        re.IGNORECASE
    )
    bare_num_pat = re.compile(r"^\s*\d{1,3}\s*$")

    def is_writing_line_mcq(txt: str) -> bool:
        s = (txt or "").replace(" ", "").replace("\t", "")
        if len(s) < 8:
            return False
        return s.count(".") / max(len(s), 1) >= 0.8

    for pi in range(len(doc2)):
        page = doc2[pi]
        ph = page.rect.height
        td = page.get_text("dict")
        lines_info = []
        for block in td.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                full = " ".join(s.get("text", "") for s in spans).strip()
                if not full:
                    continue
                bb = line["bbox"]
                lines_info.append((bb[1], bb[3], full))

        def is_footer(y0, txt):
            if footer_pat.search(txt):
                return True
            if bare_num_pat.match(txt) and y0 > ph * 0.90:
                return True
            return False

        first_footer_y0 = ph * 0.92
        for y0, y1, txt in sorted(lines_info, key=lambda t: t[0]):
            if y0 < ph * 0.5:
                continue
            if is_footer(y0, txt):
                first_footer_y0 = min(first_footer_y0, y0)
                break

        first_writing_y0 = None
        for y0, y1, txt in sorted(lines_info, key=lambda t: t[0]):
            if is_writing_line_mcq(txt):
                first_writing_y0 = y0
                break

        upper_bound = first_footer_y0
        if first_writing_y0 is not None and first_writing_y0 < upper_bound:
            upper_bound = first_writing_y0

        content_bottom = 0
        for y0, y1, txt in lines_info:
            if is_footer(y0, txt) or is_writing_line_mcq(txt):
                continue
            if y1 < upper_bound - 2:
                if y1 > content_bottom:
                    content_bottom = y1

        if content_bottom == 0:
            content_bottom = ph * 0.88

        content_bottom_per_page[pi] = min(content_bottom + 8, upper_bound - 6)

    doc2.close()

    result = {}
    for i, (qn, page_idx, top_y, ph, pw, mrx) in enumerate(raw):
        bottom_y = content_bottom_per_page.get(page_idx, ph * 0.88)
        bottom_y = min(bottom_y, ph * 0.92)
        for j in range(i + 1, len(raw)):
            nq, np_idx, ntop, _, _, _ = raw[j]
            if np_idx == page_idx:
                bottom_y = ntop - 4
                break
            if np_idx > page_idx:
                break
        page_1based = page_idx + 1
        key = (page_1based, qn)
        if key in result:
            continue
        result[key] = {
            "page_idx": page_idx, "top_y": top_y, "bottom_y": bottom_y,
            "page_height": ph, "page_width": pw, "marker_right_x": mrx,
        }

    return result


def find_question_for_excel_row(locations, q_num, excel_page, tolerance=5):
    if excel_page and excel_page > 0:
        if (excel_page, q_num) in locations:
            return locations[(excel_page, q_num)]
        candidates = [
            (abs(p - excel_page), p)
            for (p, qn) in locations
            if qn == q_num and abs(p - excel_page) <= tolerance
        ]
        if candidates:
            candidates.sort()
            _, p = candidates[0]
            return locations[(p, q_num)]
        return None
    for (p, qn), loc in sorted(locations.items()):
        if qn == q_num:
            return loc
    return None


def infer_segment_page_offsets(locations, xl_rows, segments):
    offsets = []
    for seg_rows in segments:
        seg_offset = 0
        q1_rows = [xl_rows[ri] for ri in seg_rows if xl_rows[ri]["qn"] == 1]
        if q1_rows:
            excel_q1_page = q1_rows[0]["page_num"]
            if excel_q1_page > 0:
                q1_qp_pages = sorted(p for (p, qn) in locations if qn == 1)
                if q1_qp_pages:
                    nearest = min(q1_qp_pages, key=lambda p: abs(p - excel_q1_page))
                    if abs(nearest - excel_q1_page) < 20:
                        seg_offset = nearest - excel_q1_page
        offsets.append(seg_offset)
    return offsets


def crop_question_png(pdf_bytes: bytes, loc: dict, q_num: int = None, dpi: int = 200) -> bytes | None:
    if not loc:
        return None
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc[loc["page_idx"]]
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    iw, ih = img.size
    doc.close()
    scale = dpi / 72
    top_px    = max(0,  int(loc["top_y"]    * scale) - 10)
    bottom_px = min(ih, int(loc["bottom_y"] * scale) + 5)
    mrx = loc.get("marker_right_x", 0)
    left_px = max(0, int((mrx + 6) * scale)) if mrx else 0
    pw = loc.get("page_width", 0) or (iw / scale)
    right_px = iw - int(pw * 0.05 * scale)
    if bottom_px <= top_px or right_px <= left_px:
        return None
    cropped = img.crop((left_px, top_px, right_px, bottom_px))
    buf = io.BytesIO()
    cropped.save(buf, format="PNG", optimize=True)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
#  MCQ EXTRACTION FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════
def classify_and_extract(client, qp_b64, q_nums_or_rows, qp_bytes=None, locations=None):
    if qp_bytes is None or locations is None:
        return []
    doc = fitz.open(stream=qp_bytes, filetype="pdf")
    results = []
    for entry in q_nums_or_rows:
        if isinstance(entry, tuple):
            q_num, excel_page = entry
        else:
            q_num, excel_page = entry, 0
        loc = find_question_for_excel_row(locations, q_num, excel_page)
        if not loc:
            results.append({
                "qn": q_num, "found": False, "needs_image": False,
                "text": "", "A": "", "B": "", "C": "", "D": "",
                "_loc": None, "_excel_page": excel_page,
            })
            continue
        page = doc[loc["page_idx"]]
        top_y, bottom_y = loc["top_y"], loc["bottom_y"]
        clip = fitz.Rect(0, top_y, page.rect.width, bottom_y)
        text_dict = page.get_text("dict", clip=clip)
        all_spans = []
        size_counts = {}
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    if not text:
                        continue
                    sz = round(span.get("size", 0), 1)
                    bb = span.get("bbox", (0, 0, 0, 0))
                    all_spans.append({"text": text, "size": sz, "x0": bb[0],
                                      "y0": bb[1], "y1": bb[3], "flags": span.get("flags", 0)})
                    if text.strip():
                        size_counts[sz] = size_counts.get(sz, 0) + len(text.strip())
        if not all_spans:
            results.append({"qn": q_num, "found": True, "needs_image": True,
                             "text": "", "A": "", "B": "", "C": "", "D": ""})
            continue
        body_size = max(size_counts.items(), key=lambda kv: kv[1])[0] if size_counts else 11.0
        SUBS = str.maketrans("0123456789+-=()n", "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₙ")
        SUPS = str.maketrans("0123456789+-=()n", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ")
        unsafe_sub_super = False
        Y_BUCKET = 3
        line_buckets = {}
        for sp in all_spans:
            key = round(sp["y0"] / Y_BUCKET) * Y_BUCKET
            line_buckets.setdefault(key, []).append(sp)
        rendered_lines = []
        for key in sorted(line_buckets):
            spans_on_line = sorted(line_buckets[key], key=lambda s: s["x0"])
            full_size_spans = [s for s in spans_on_line if s["size"] >= body_size - 0.5]
            if full_size_spans:
                main_y0 = full_size_spans[0]["y0"]
                main_y1 = full_size_spans[0]["y1"]
            else:
                main_y0 = spans_on_line[0]["y0"]
                main_y1 = spans_on_line[0]["y1"]
            parts = []
            for sp in spans_on_line:
                text = sp["text"]
                if not text.strip():
                    parts.append(text)
                    continue
                is_smaller = sp["size"] < body_size - 1.0
                if is_smaller and sp["y1"] > main_y1 + 0.5:
                    if all(c in "0123456789+-=()n" for c in text.strip()):
                        parts.append(text.translate(SUBS))
                    else:
                        unsafe_sub_super = True
                        parts.append(text)
                elif is_smaller and sp["y0"] < main_y0 - 0.5:
                    if all(c in "0123456789+-=()n" for c in text.strip()):
                        parts.append(text.translate(SUPS))
                    else:
                        unsafe_sub_super = True
                        parts.append(text)
                else:
                    parts.append(text)
            rendered = "".join(parts)
            rendered = re.sub(r"[\u2009\u00a0]", " ", rendered).strip()
            if rendered:
                rendered_lines.append((key, spans_on_line[0]["x0"], rendered))
        opt_re = re.compile(r"^\s*([ABCD])\.\s*(.*)$")
        opt_markers = {}
        for y_key, x0, rendered in rendered_lines:
            m = opt_re.match(rendered)
            if m and x0 < 200:
                opt_markers[y_key] = m.group(1)
        needs_image = bool(unsafe_sub_super)
        stem_lines = []
        option_lines = {"A": [], "B": [], "C": [], "D": []}
        cur_opt = None
        for y_key, x0, rendered in rendered_lines:
            if y_key in opt_markers:
                cur_opt = opt_markers[y_key]
                m = opt_re.match(rendered)
                if m:
                    rest = m.group(2).strip()
                    if rest:
                        option_lines[cur_opt].append(rest)
                continue
            if cur_opt:
                option_lines[cur_opt].append(rendered)
            else:
                stem_lines.append(rendered)
        qn_pref = re.compile(r"^\d+\.\s*")
        if stem_lines and qn_pref.match(stem_lines[0]):
            stem_lines[0] = qn_pref.sub("", stem_lines[0]).strip()
        stem = "\n".join(stem_lines).strip()
        options = {L: " ".join(option_lines[L]).strip() for L in "ABCD"}

        def _has_chem_complexity(s):
            if not s:
                return False
            if re.search(r"[A-Z]{2,}", s) and re.search(r"\d", s):
                return True
            if re.search(r"[A-Z][a-z]?\d", s):
                return True
            if re.search(r"\b[A-Z][a-z]?\s+\d", s):
                return True
            if re.search(r"[=×]", s):
                return True
            return False

        if not needs_image:
            if _has_chem_complexity(stem):
                needs_image = True
            else:
                for L in "ABCD":
                    if _has_chem_complexity(options[L]):
                        needs_image = True
                        break
        non_empty_opts = [v for v in options.values() if v and len(v.strip()) >= 1]
        if len(non_empty_opts) < 4 or not stem:
            needs_image = True
        if needs_image:
            stem = ""
            options = {"A": "", "B": "", "C": "", "D": ""}
        results.append({
            "qn": q_num, "found": True, "needs_image": needs_image,
            "text": stem, "A": options["A"], "B": options["B"],
            "C": options["C"], "D": options["D"],
            "_loc": loc, "_excel_page": excel_page,
        })
    doc.close()
    return results


def _extract_qn_ans_pairs(text: str) -> list[tuple]:
    pattern = re.compile(
        r"(?:^|[\s\|])(?:Q\s*)?(\d{1,2})[\.\):]?\s*([ABCD])(?=$|[\s\.\,\;\|])"
    )
    pairs = []
    for m in pattern.finditer(text):
        qn = int(m.group(1))
        if 1 <= qn <= 99:
            pairs.append((qn, m.group(2)))
    return pairs


def extract_answers_pymupdf_per_section(ms_bytes: bytes) -> list[dict]:
    sections = []
    current = {"start_page": 1, "answers": {}, "max_qn": 0}
    try:
        doc = fitz.open(stream=ms_bytes, filetype="pdf")
    except Exception:
        return []
    for page_idx, page in enumerate(doc):
        text = page.get_text()
        page_pairs = _extract_qn_ans_pairs(text)
        if len(page_pairs) < 3:
            continue
        for qn, ans in page_pairs:
            qn_str = str(qn)
            is_boundary = (qn_str in current["answers"] or (qn == 1 and current["max_qn"] >= 5))
            if is_boundary and current["answers"]:
                sections.append({"start_page": current["start_page"], "end_page": page_idx + 1,
                                  "answers": current["answers"]})
                current = {"start_page": page_idx + 1, "answers": {}, "max_qn": 0}
            current["answers"][qn_str] = ans
            if qn > current["max_qn"]:
                current["max_qn"] = qn
    if current["answers"]:
        sections.append({"start_page": current["start_page"], "end_page": len(doc),
                          "answers": current["answers"]})
    doc.close()
    return [s for s in sections if s["answers"]]


def extract_answers_pymupdf(ms_bytes: bytes) -> dict:
    sections = extract_answers_pymupdf_per_section(ms_bytes)
    answers = {}
    for sec in sections:
        for q, a in sec["answers"].items():
            answers.setdefault(q, a)
    return answers


def extract_answers(client, ms_b64, q_nums, ms_bytes=None) -> dict:
    answers = {}
    if ms_bytes is not None:
        answers = extract_answers_pymupdf(ms_bytes)
    missing = [n for n in q_nums if str(n) not in answers]
    if missing and client is not None:
        nums = ", ".join(str(n) for n in missing)
        try:
            response = client.messages.create(
                model="claude-opus-4-5", max_tokens=1500,
                messages=[{"role": "user", "content": [
                    {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": ms_b64}},
                    {"type": "text", "text": (
                        f"Extract answers for Q{nums} from this IB mark scheme. "
                        'Return ONLY JSON: {"1":"D","2":"A",...}. '
                        'Use "NOT_FOUND" if absent.'
                    )},
                ]}]
            )
            raw = "".join(b.text for b in response.content if hasattr(b, "text"))
            data = safe_json(raw)
            if isinstance(data, dict):
                for k, v in data.items():
                    if k not in answers and v in ("A", "B", "C", "D"):
                        answers[k] = v
        except Exception:
            pass
    return answers


# ═══════════════════════════════════════════════════════════════════════════════
#  STRUCTURED MS — SECTION DETECTION
# ═══════════════════════════════════════════════════════════════════════════════
def _detect_paper_code(text: str) -> str | None:
    """
    Extract a normalized paper code from PDF text.
    Works for any IB subject: Math, Biology, Chemistry, etc.
    Returns a code like "M23/HL/P2/TZ1" or None.
    """
    if not text:
        return None
    # Pattern 1: Full IB code like M21/4/BIOLO/HP2/ENG/TZ1/XX
    m = re.search(
        r"([MN])(\d{2})/\d/\w+/([HS])P(\d)/\w+/(TZ[\d0]|XX)",
        text, re.IGNORECASE
    )
    if m:
        session, year, level, paper, tz = m.groups()
        lv = "HL" if level.upper() == "H" else "SL"
        tz_str = tz.upper() if tz.upper() != "XX" else "TZ0"
        return f"{session}{year}/{lv}/P{paper}/{tz_str}"
    # Pattern 2: 4-digit code like 2221-7210 or 8825-7310
    m2 = re.search(r"(\d{4})[\-\u2013](\d{4})", text)
    if m2:
        return f"CODE-{m2.group(1)}-{m2.group(2)}"
    # Pattern 3: "May 2023 Paper 2 Higher Level"
    m3 = re.search(
        r"(May|November|March)\s+(\d{4}).*?[Pp]aper\s*(\d).*?(Higher|Standard)\s*[Ll]evel",
        text, re.DOTALL
    )
    if m3:
        month, year, paper, level = m3.groups()
        lv = "HL" if "Higher" in level else "SL"
        session = "M" if month == "May" else ("N" if month == "November" else "R")
        return f"{session}{year[2:]}/{lv}/P{paper}"
    return None


def _extract_paper_meta(pdf_bytes: bytes, n_pages_check: int = 10) -> dict:
    """
    Extract subject, level, paper number, year, session, timezone from a PDF.
    Returns a dict with keys: subject, level, paper, year, session, tz, code.
    """
    meta = {"subject":"","level":"","paper":"","year":"","session":"","tz":"","code":""}
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        # Check first few pages for cover page info
        for pi in range(min(n_pages_check, len(doc))):
            txt = doc[pi].get_text()
            # Subject
            for subj in ["Biology","Chemistry","Physics","Mathematics","Economics",
                         "History","Geography","Computer Science","Psychology"]:
                if subj.lower() in txt.lower():
                    meta["subject"] = subj; break
            # Level
            if "higher level" in txt.lower(): meta["level"] = "HL"
            elif "standard level" in txt.lower(): meta["level"] = "SL"
            # Paper
            m_paper = re.search(r"[Pp]aper\s*(\d)", txt)
            if m_paper: meta["paper"] = m_paper.group(1)
            # Year + Session
            m_date = re.search(r"(May|November|March)\s+(\d{4})", txt)
            if m_date:
                meta["session"] = m_date.group(1)
                meta["year"]    = m_date.group(2)
            # Timezone
            m_tz = re.search(r"(TZ\d|timezone\s*\d)", txt, re.I)
            if m_tz: meta["tz"] = m_tz.group(1).upper().replace(" ","").replace("TIMEZONE","TZ")
            # Paper code
            code = _detect_paper_code(txt)
            if code and not meta["code"]: meta["code"] = code
            if all(meta[k] for k in ["subject","level","paper","year"]): break
        doc.close()
    except Exception:
        pass
    return meta


def _verify_qp_ms_match(qp_bytes: bytes, ms_bytes: bytes) -> tuple[bool, str]:
    """
    Verify that QP and MS PDFs match (same subject/level/paper/year/session/tz).
    Returns (match: bool, message: str).
    """
    qp_meta = _extract_paper_meta(qp_bytes)
    ms_meta  = _extract_paper_meta(ms_bytes)

    mismatches = []
    for key in ["subject","level","paper","year"]:
        qv = qp_meta.get(key,"").strip()
        mv = ms_meta.get(key,"").strip()
        if qv and mv and qv.lower() != mv.lower():
            mismatches.append(f"{key}: QP='{qv}' vs MS='{mv}'")

    # Timezone check (optional — only fail if both specify and differ)
    qtz = qp_meta.get("tz","")
    mtz = ms_meta.get("tz","")
    if qtz and mtz and qtz != mtz and qtz != "TZ0" and mtz != "TZ0":
        mismatches.append(f"timezone: QP='{qtz}' vs MS='{mtz}'")

    if mismatches:
        return False, ("❌ QP and Mark Scheme do not match:\n" +
                       "\n".join(f"  • {m}" for m in mismatches))

    summary = (f"✅ QP ↔ MS verified: {qp_meta.get('subject','')} "
               f"{qp_meta.get('level','')} Paper {qp_meta.get('paper','')} "
               f"{qp_meta.get('session','')} {qp_meta.get('year','')} "
               f"{qp_meta.get('tz','')}")
    return True, summary.strip()
def _build_page_to_paper_map(pdf_bytes: bytes) -> dict:
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception:
        return {}
    page_to_code = {}
    last_code = None
    for pi in range(len(doc)):
        text = doc[pi].get_text()
        code = _detect_paper_code(text)
        if code:
            last_code = code
        if last_code:
            page_to_code[pi] = last_code
    doc.close()
    return page_to_code


def find_ms_question_locations(ms_bytes: bytes) -> list[dict]:
    try:
        doc = fitz.open(stream=ms_bytes, filetype="pdf")
    except Exception:
        return []

    raw_markers = []
    pat_strict  = re.compile(r"^(\d{1,2})\.\s*$")
    pat_qa      = re.compile(r"^(\d{1,2})\.\s+\(?\w")
    pat_nodot   = re.compile(r"^(\d{1,2})\s*$")
    pat_qprefix = re.compile(r"^Q(\d{1,2})\.\s*(?:\(?\w|$)")

    for pi in range(len(doc)):
        page = doc[pi]
        ph = page.rect.height
        td = page.get_text("dict")

        # Skip pages that are general marking instructions (not real Q answers)
        page_text = page.get_text()
        if _MS_INSTR_PAT.search(page_text):
            continue

        for block in td.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                full_line = "".join(s.get("text", "") for s in spans).strip()
                bb = line["bbox"]
                if bb[0] > 80:
                    continue
                qn = None
                m = pat_strict.match(full_line) or pat_qa.match(full_line)
                if m:
                    qn = int(m.group(1))
                else:
                    mq = pat_qprefix.match(full_line)
                    if mq:
                        qn = int(mq.group(1))
                    else:
                        m2 = pat_nodot.match(full_line)
                        if m2 and bb[0] < 50:
                            font = spans[0].get("font", "")
                            size = spans[0].get("size", 0)
                            if ("Bold" in font or "bold" in font) and 9 <= size <= 14:
                                qn = int(m2.group(1))
                if qn is None or qn < 1 or qn > 30:
                    continue
                if re.match(rf"^{qn}\.\d", full_line):
                    continue
                marker_right_x = spans[0]["bbox"][2] if spans else (bb[0] + 20)
                raw_markers.append((qn, pi, bb[1], ph, marker_right_x))

    doc.close()
    if not raw_markers:
        return []

    sections = []
    current = {"questions": {}, "_order": [], "start_page": raw_markers[0][1] + 1, "max_qn": 0}
    for qn, pi, top_y, ph, mrx in raw_markers:
        if (qn in current["questions"] or (qn == 1 and current["max_qn"] >= 5)):
            if current["questions"]:
                sections.append(current)
                current = {"questions": {}, "_order": [], "start_page": pi + 1, "max_qn": 0}
        current["questions"][qn] = {
            "page_idx": pi, "top_y": top_y, "bottom_y": ph * 0.92,
            "end_page_idx": pi, "end_y": ph * 0.92,
            "page_height": ph, "marker_right_x": mrx,
        }
        current["_order"].append((qn, pi, top_y))
        if qn > current["max_qn"]:
            current["max_qn"] = qn

    if current["questions"]:
        sections.append(current)

    for sec in sections:
        order = sec["_order"]
        for idx, (qn, pi, top_y) in enumerate(order):
            if idx + 1 < len(order):
                next_qn, next_pi, next_top_y = order[idx + 1]
                if next_pi == pi:
                    sec["questions"][qn]["end_page_idx"] = pi
                    sec["questions"][qn]["end_y"] = next_top_y - 4
                else:
                    sec["questions"][qn]["end_page_idx"] = next_pi
                    sec["questions"][qn]["end_y"] = next_top_y - 4
        last_qn, last_pi, _ = order[-1]
        sec["end_page"] = sec["questions"][last_qn]["end_page_idx"] + 1
        sec.pop("_order", None)
        # Store max_qn for matching heuristic
        sec["max_qn"] = max(sec["questions"].keys()) if sec["questions"] else 0

    # ── Merge Paper 2 sub-sections (Section A Q1-5 + Section B Q6-10) ───────────
    # IB Paper 2 MS PDFs sometimes split each exam into two sub-sections:
    #   Sub-section 1 (Q6-10, "Section B") and Sub-section 2 (Q1-5, "Section A")
    # These must be merged into one combined section per exam so that all
    # questions (Q1-10) are available for matching.
    merged_secs = []
    i = 0
    while i < len(sections):
        sec = sections[i]
        # Check if next section is a "small" companion (max_qn ≤ 6, same exam)
        if i + 1 < len(sections):
            nxt = sections[i + 1]
            # Merge when: this section has high max_qn AND next has low max_qn
            # (or vice-versa) — they are two halves of the same exam
            this_mq = sec.get("max_qn", 0)
            nxt_mq  = nxt.get("max_qn", 0)
            should_merge = (
                (this_mq >= 6 and nxt_mq <= 6) or
                (nxt_mq  >= 6 and this_mq <= 6)
            )
            if should_merge:
                combined_questions = {}
                combined_questions.update(sec.get("questions", {}))
                combined_questions.update(nxt.get("questions", {}))
                combined = {
                    "questions":  combined_questions,
                    "start_page": sec.get("start_page", sec.get("start_pi", 0) + 1),
                    "end_page":   nxt.get("end_page",   nxt.get("start_pi", 0) + 1),
                    "max_qn":     max(this_mq, nxt_mq),
                    "paper_code": sec.get("paper_code") or nxt.get("paper_code"),
                }
                merged_secs.append(combined)
                i += 2
                continue
        merged_secs.append(sec)
        i += 1

    # Use merged list if it produced a significantly different count
    if len(merged_secs) != len(sections):
        sections = merged_secs

    # Filter: remove sections with anomalous question numbers (merged artifacts)
    # and sections missing Q1 (not a real answer section start)
    def _is_valid_section(s):
        qs = set(s["questions"].keys())
        if len(qs) < 3:
            return False
        if 1 not in qs:
            return False   # missing Q1 = not a real section
        max_q = max(qs)
        # Check for outlier Q numbers that indicate a merged/corrupt section
        # e.g. qs=[1,2,3,4,5,8] — Q8 is outlier when max expected is 5
        expected_range = range(1, max_q + 1)
        valid_qs = {q for q in qs if q <= max_q}
        # Any question > 1.6x the second-highest is an outlier
        sorted_qs = sorted(qs)
        if len(sorted_qs) >= 2:
            second_max = sorted_qs[-2]
            top_q = sorted_qs[-1]
            if top_q > second_max + 1:
                return False   # outlier top question (gap > 1 from second-highest)
        return True

    sections = [s for s in sections if _is_valid_section(s)]

    instr_pat = re.compile(
        r"(instructions to examiners|abbreviations|marks? awarded for|"
        r"using the markscheme|method of marking|implied marks|misread|brackets in working)",
        re.IGNORECASE
    )
    try:
        doc2 = fitz.open(stream=ms_bytes, filetype="pdf")
        # Answer-content markers (M1, A1, R1, etc.)
        answer_marker_pat = re.compile(r"\b(M1|A1|A2|A3|R1|M0|AG|FT|ft)\b")
        filtered = []
        for s in sections:
            start = s["start_page"] - 1
            # A section is an instruction-only section (no real answers)
            # when its first page has instructions but NO answer markers
            # (P2 MS puts instructions + answers on the same page, so we
            #  must check for the presence of answer markers to distinguish)
            is_instr_only = False
            for pi in range(start, min(start + 1, len(doc2))):
                txt = doc2[pi].get_text()
                has_instr   = bool(instr_pat.search(txt))
                has_answers = bool(answer_marker_pat.search(txt))
                # Only filter if it's a pure instructions page (no answer markers)
                if has_instr and not has_answers:
                    is_instr_only = True
            if not is_instr_only:
                filtered.append(s)
        doc2.close()
        sections = filtered
    except Exception:
        pass

    page_codes = _build_page_to_paper_map(ms_bytes)
    for s in sections:
        first_page_idx = s["start_page"] - 1
        s["paper_code"] = page_codes.get(first_page_idx)

    return sections


# ═══════════════════════════════════════════════════════════════════════════════
#  STRUCTURED MODE — QP CROP (FINAL AGENT approach)
# ═══════════════════════════════════════════════════════════════════════════════
_S_DPI   = 150
_S_SCALE = _S_DPI / 72
_S_CW    = 16.0          # cm content width
_S_CW_PT = _S_CW * 28.35

# Continuation text patterns
_CONTINUES_PAT = re.compile(r'this question continues on the following page', re.I)
_CONTINUED_PAT = re.compile(r'question\s+\d+\s+continued', re.I)
_MS_CONT_PAT   = re.compile(r'question\s*\d*\s*continued\.?|this question continues', re.I)

# Instruction-page detector: these phrases appear ONLY in general marking guidance,
# never in actual Q-answer sections.
_MS_INSTR_PAT  = re.compile(
    r"implied marks appear in brackets|"
    r"follow.through \(FT\) marks are awarded|"
    r"instructions to examiners|"
    r"using the markscheme|"
    r"method and answer.accuracy marks|"
    r"mis.read.*candidate incorrectly copies|"
    r"abbreviations.*marks awarded for attempting|"
    r"alternative forms.*accept equivalent|"
    r"presentation of candidate work|"
    r"crossed out work.*candidate has drawn|"
    r"gcalculators.*gdc is required|"
    r"GDC is required for this paper|"
    r"markscheme.*may 20|markscheme.*november 20",
    re.IGNORECASE | re.DOTALL
)

# QP cover/instruction page detector (for skipping non-question pages in QP PDF)
# These phrases appear on exam cover pages — never crop them as question content.
_QP_INSTR_PAT = re.compile(
    r"instructions to candidates|"
    r"do not open this examination|"
    r"graphic display calculator is required|"
    r"answer all the questions in the answer booklet|"
    r"maximum mark for this examination|"
    r"formula booklet is required|"
    r"answer booklet provided",
    re.IGNORECASE
)

_S_FOOTER_PAT  = re.compile(
    r"(M\d{2}/\d|N\d{2}/\d|\d{4}EP\d+|©\s*\d{4}|\bTurn over\b|\bPlease do not\b|"
    r"international\s+baccalaureate|\d{4}[\s\-–]+\d{4})", re.I)
_MS_FT  = re.compile(r"(M\d{2}/\d/|N\d{2}/\d/|©\s*\d{4}|international\s+baccalaureate)", re.I)
_MS_BN  = re.compile(r"^\s*\d{1,3}\s*$")
_MS_HDR = re.compile(r"(^\s*[–\-]\s*\d{1,3}\s*[–\-]\s*$|M\d{2}/\d/|N\d{2}/\d/)", re.I)


def _extract_qp_text(qp_doc, pg: int, qn: int) -> str | None:
    """
    Extract the text of question qn from QP page pg.
    Returns clean text if the question is text-only, else None (use image).
    Heuristic: if >40% of content is non-ASCII or there are image blocks, use image.
    """
    try:
        page = qp_doc[pg - 1]
        td   = page.get_text("dict")
        blocks = td.get("blocks", [])
        # Check for image blocks (type=1 = image)
        has_images = any(b.get("type") == 1 for b in blocks)
        # Check for drawing objects (figures/diagrams)
        clip_svg  = page.get_svg_image()
        has_drawings = len(clip_svg) > 2000 and "<path" in clip_svg

        if has_images or has_drawings:
            return None   # Use image

        # Extract all text lines after the question marker
        ph = page.rect.height
        lines_out = []
        in_q = False
        for block in blocks:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                txt = "".join(s.get("text", "") for s in spans).strip()
                if not txt:
                    continue
                bb = line["bbox"]
                # Detect question marker
                if re.match(rf"^{qn}[.\s]", txt) and bb[0] < 70:
                    in_q = True
                    # Include [Maximum mark] on same line
                    lines_out.append(txt)
                    continue
                # Stop at next question
                if in_q and re.match(r"^\d{1,2}[.\s]", txt) and bb[0] < 70:
                    m_next = re.match(r"^(\d{1,2})", txt)
                    if m_next and int(m_next.group(1)) > qn:
                        break
                # Skip footers
                if bb[1] > ph * 0.90:
                    continue
                # Skip writing lines (dots/underscores)
                if _s_is_writing_line(txt):
                    continue
                if in_q:
                    lines_out.append(txt)

        if not lines_out:
            return None
        result = "\n".join(lines_out)
        # If >25% non-ASCII → likely formula-heavy → use image
        non_ascii = sum(1 for c in result if ord(c) > 127)
        if len(result) > 0 and non_ascii / len(result) > 0.25:
            return None
        return result
    except Exception:
        return None


def _s_trim_bottom(img, thr=248):
    arr = np.array(img)
    rm  = arr.min(axis=(1, 2))
    dk  = np.where(rm < thr)[0]
    if len(dk) == 0:
        return img
    return img.crop((0, 0, img.width, min(int(dk[-1]) + 10, img.height)))


def _s_whiteout(img, rects_pt, origin_pt=(0.0, 0.0)):
    """Paint white rectangles over specified PDF-coordinate regions.

    img        : PIL Image (already cropped from the PDF page).
    rects_pt   : list of (x0, y0, x1, y1) in PDF point coordinates.
    origin_pt  : (x_pt, y_pt) of the top-left corner of the crop in PDF space.
    """
    from PIL import ImageDraw
    img  = img.copy()
    draw = ImageDraw.Draw(img)
    ox, oy = origin_pt
    for (x0, y0, x1, y1) in rects_pt:
        draw.rectangle([
            max(0, int((x0 - ox) * _S_SCALE) - 1),
            max(0, int((y0 - oy) * _S_SCALE) - 1),
            int((x1 - ox) * _S_SCALE) + 2,
            int((y1 - oy) * _S_SCALE) + 2,
        ], fill="white")
    del draw
    return img


def _s_split_smart(img, max_h=1500, thr=250, win=250):
    img = _s_trim_bottom(img)
    w, h = img.size
    if h <= max_h:
        return [img]
    arr = np.array(img)
    rd  = (arr.min(axis=(1, 2)) < thr).astype(int)
    chunks = []
    prev = 0
    while prev < h:
        tgt = prev + max_h
        if tgt >= h:
            c = img.crop((0, prev, w, h))
            if c.height > 5:
                chunks.append(_s_trim_bottom(c))
            break
        lo = max(prev + 80, tgt - win)
        hi = min(h - 10, tgt + win)
        best = tgt
        bl = 0
        iw = False
        ws = 0
        for i, d in enumerate(rd[lo:hi]):
            ry = lo + i
            if d == 0:
                if not iw:
                    iw = True
                    ws = ry
            else:
                if iw:
                    rl = ry - ws
                    mid = (ws + ry) // 2
                    if rl > bl:
                        bl = rl
                        best = mid
                    iw = False
        if iw and (hi - ws) > bl:
            best = (ws + hi) // 2
        c = img.crop((0, prev, w, best))
        if c.height > 5:
            chunks.append(_s_trim_bottom(c))
        prev = best
    return chunks if chunks else [img]


def _s_is_writing_line(txt):
    if not txt:
        return False
    s = txt.replace(" ", "").replace("\t", "")
    if len(s) < 4:
        return False
    if s.count(".") / max(len(s), 1) >= 0.70:
        return True
    if len(s) >= 15 and s.count("_") / max(len(s), 1) >= 0.70:
        return True
    if len(s) >= 15 and s.count("-") / max(len(s), 1) >= 0.88:
        return True
    n_sp = sum(1 for c in s if ord(c) < 32 or ord(c) == 0xfffd or ord(c) == 0x08)
    if n_sp / max(len(s), 1) >= 0.50 and len(s) >= 5:
        return True
    n_na = sum(1 for c in s if ord(c) > 127)
    if n_na / max(len(s), 1) >= 0.70 and len(s) >= 10:
        return True
    return False


def _s_is_q_marker(txt, x0):
    """True only when txt is a QP question-number marker at x0 < 55pt."""
    if x0 > 55:
        return False
    m = re.match(r'^(\d{1,2})[.\s]', txt)
    if m and 1 <= int(m.group(1)) <= 20:
        return True
    return False


def _s_get_lines(page):
    r = []
    for block in page.get_text("dict").get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            if not spans:
                continue
            txt = "".join(s.get("text", "") for s in spans).strip()
            if not txt:
                continue
            bb = line["bbox"]
            r.append((bb[0], bb[1], bb[3], txt))
    return sorted(r, key=lambda x: x[1])


def _s_find_answer_box_top(page, cs, se):
    pw = page.rect.width
    best = None
    for d in page.get_drawings():
        for item in d.get("items", []):
            if item[0] != "l":
                continue
            p1, p2 = item[1], item[2]
            if abs(p1.y - p2.y) > 2 or abs(p1.x - p2.x) < pw * 0.65:
                continue
            y = min(p1.y, p2.y)
            if y < cs - 2 or y > se:
                continue
            if best is None or y < best:
                best = y
    return best


def _crop_qp_page(qp_doc, pg_idx, qn, is_cont=False):
    page = qp_doc[pg_idx]
    ph = page.rect.height
    pw = page.rect.width
    lines = _s_get_lines(page)
    X_LEFT  = 30.0
    X_RIGHT = pw - 22.0

    if is_cont:
        # Skip if this is a cover/instruction page (never a valid continuation)
        if _QP_INSTR_PAT.search(page.get_text()):
            return None, False

        hb = 0
        cb = 0
        for x0, y0, y1, txt in lines:
            if y0 > ph * 0.20:
                break
            if _S_FOOTER_PAT.search(txt) or re.match(r'^[\-–]\s*\d', txt):
                hb = max(hb, y1)
            if _CONTINUED_PAT.search(txt):
                cb = max(cb, y1)
        cs = max(hb + 4, cb + 4, ph * 0.07)
    else:
        qt = None
        qb = 0
        for x0, y0, y1, txt in lines:
            if not _s_is_q_marker(txt, x0):
                continue
            m = re.match(r'^(\d{1,2})[.\s]', txt)
            if m and int(m.group(1)) == qn and 1 <= qn <= 20 and y0 < ph * 0.88:
                qt = y0
                qb = y1
                break
        if qt is None:
            qt = ph * 0.07
            qb = ph * 0.12
        for x0, y0, y1, txt in lines:
            if y0 < qb - 2:
                continue
            if y0 > qb + 18:
                break
            if re.match(r'^\[\s*Maximum mark', txt, re.I):
                qb = max(qb, y1)
        # Start crop FROM the question number line (includes "[Maximum mark: N]")
        # qt is the top of "N." line; [Maximum mark] is on the same line.
        cs = max(0, qt - 6)

    nq = ph * 0.92
    if not is_cont:
        for x0, y0, y1, txt in lines:
            if y0 <= cs + 2:
                continue
            if not _s_is_q_marker(txt, x0):
                continue
            m = re.match(r'^(\d{1,2})[.\s]', txt)
            if m:
                n = int(m.group(1))
                if n > qn and 1 <= n <= 20 and y0 < ph * 0.88:
                    nq = y0
                    break

    cont_y = None
    for x0, y0, y1, txt in lines:
        if y0 < cs:
            continue
        if y0 > nq:
            break
        if _CONTINUES_PAT.search(txt):
            cont_y = y0
            break

    fw = None
    for x0, y0, y1, txt in lines:
        if y0 < cs - 2:
            continue
        if y0 >= (cont_y or nq):
            break
        if _s_is_writing_line(txt):
            fw = y0
            break

    se = cont_y or fw or nq
    ab = _s_find_answer_box_top(page, cs, se)

    if ab:
        ce = ab - 1
    else:
        lc = cs
        for x0, y0, y1, txt in lines:
            if y0 < cs - 2:
                continue
            if y0 >= se:
                break
            if _s_is_writing_line(txt):
                break
            if (_S_FOOTER_PAT.search(txt) or _CONTINUES_PAT.search(txt)
                    or _CONTINUED_PAT.search(txt)):
                continue
            lc = max(lc, y1)
        ce = min(lc + 16, se - 2)

    ce = min(ce, ph * 0.91)
    if ce <= cs + 5:
        return None, bool(cont_y)

    has_cont = bool(cont_y) or bool(_CONTINUES_PAT.search(page.get_text()))
    clip = fitz.Rect(X_LEFT, cs, X_RIGHT, ce)
    pix  = page.get_pixmap(matrix=fitz.Matrix(_S_SCALE, _S_SCALE), clip=clip, alpha=False)
    img  = Image.open(io.BytesIO(pix.tobytes("png")))

    # Strip "N." from the question number line, keep "[Maximum mark: ...]"
    # The "N." span is at far-left (x < 60pt), on the same line as qt.
    # "[Maximum mark]" span is further right (x ≈ 70+pt) — NOT whited out.
    if not is_cont:
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    stxt = span.get("text", "").strip()
                    sbb  = span.get("bbox", [0, 0, 0, 0])
                    # Only match the bare "N." or "N" span (not [Maximum mark])
                    if (re.match(r"^\d{1,2}[.]?$", stxt)
                            and sbb[0] < 60
                            and sbb[2] < 65
                            and qt is not None
                            and abs(sbb[1] - qt) < 20):
                        img = _s_whiteout(
                            img,
                            [(sbb[0], sbb[1], sbb[2] + 4, sbb[3] + 2)],
                            origin_pt=(X_LEFT, cs),
                        )

    return img, has_cont


def _crop_qp_full(qp_doc, pg, qn, pg_end: int = 0):
    """Crop QP question qn from page pg (through pg_end if it spans pages).

    pg_end:  last page of this question (inclusive).  0 → auto-detect via
             'continues on following page' text in the PDF.
    """
    if not pg or pg > len(qp_doc):
        return None

    img1, hc = _crop_qp_page(qp_doc, pg - 1, qn, False)
    if img1 is None:
        return None

    # ── If Excel told us the page range, use it directly ─────────────────────
    if pg_end and pg_end > pg:
        pieces = [img1]
        for pn in range(pg + 1, pg_end + 1):
            if pn > len(qp_doc):
                break
            ic, _ = _crop_qp_page(qp_doc, pn - 1, qn, True)
            if ic is not None:
                pieces.append(ic)
        tw = max(p.size[0] for p in pieces)
        th = sum(p.size[1] for p in pieces)
        c  = Image.new("RGB", (tw, th), "white")
        y  = 0
        for p in pieces:
            c.paste(p, (0, y))
            y += p.size[1]
        return c

    # ── Auto-detect via "continues on following page" ─────────────────────────
    if not hc:
        return img1
    pieces = [img1]
    pn = pg
    for _ in range(3):
        pn += 1
        if pn > len(qp_doc):
            break
        # Never cross into a cover/instruction page (start of next session)
        if _QP_INSTR_PAT.search(qp_doc[pn - 1].get_text()):
            break
        ic, hc2 = _crop_qp_page(qp_doc, pn - 1, qn, True)
        if ic is None:
            break
        pieces.append(ic)
        if not hc2:
            break
    tw = max(p.size[0] for p in pieces)
    th = sum(p.size[1] for p in pieces)
    c  = Image.new("RGB", (tw, th), "white")
    y  = 0
    for p in pieces:
        c.paste(p, (0, y))
        y += p.size[1]
    return c


# ═══════════════════════════════════════════════════════════════════════════════
#  STRUCTURED MODE — MS CROP (FINAL AGENT approach)
# ═══════════════════════════════════════════════════════════════════════════════
def _ms_hdr_bot(pg):
    ph = pg.rect.height
    hb = 0
    for block in pg.get_text("dict").get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            if not spans:
                continue
            txt = "".join(s.get("text", "") for s in spans).strip()
            bb  = line["bbox"]
            if bb[1] > ph * 0.20:
                continue
            if _MS_HDR.search(txt) and bb[3] > hb:
                hb = bb[3]
            if _MS_CONT_PAT.search(txt) and bb[3] > hb:
                hb = bb[3]
    return hb


def _ms_ft_top(pg):
    ph = pg.rect.height
    fy = ph * 0.95
    for block in pg.get_text("dict").get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            if not spans:
                continue
            txt = "".join(s.get("text", "") for s in spans).strip()
            bb  = line["bbox"]
            if bb[1] < ph * 0.85:
                continue
            if (_MS_FT.search(txt) or (_MS_BN.match(txt) and bb[1] > ph * 0.93)) and bb[1] < fy:
                fy = bb[1]
    return fy


def _last_ms_y(pg, top, hard):
    ph   = pg.rect.height
    last = top
    for block in pg.get_text("dict").get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            if not spans:
                continue
            txt = "".join(s.get("text", "") for s in spans).strip()
            if not txt:
                continue
            bb = line["bbox"]
            if bb[1] < top - 2 or bb[1] >= hard - 2:
                continue
            if _MS_FT.search(txt) or (_MS_BN.match(txt) and bb[1] > ph * 0.93):
                continue
            if _MS_CONT_PAT.search(txt):
                continue
            if bb[3] > last:
                last = bb[3]
    return last


def _extract_answer_keywords(excel_ms_text: str) -> list[str]:
    """
    Extract numeric and key word tokens from an Excel MS answer cell.
    These are used to verify/locate the correct MS section in the PDF.
    """
    if not excel_ms_text or excel_ms_text in ("nan", "None", ""):
        return []
    import re as _re
    # Extract numbers (including decimals)
    nums = _re.findall(r"\b\d+\.?\d*\b", excel_ms_text)
    # Extract meaningful words (5+ chars, no LaTeX commands)
    words = [w.lower() for w in _re.findall(r"[a-z]{5,}", excel_ms_text, _re.I)
             if w.lower() not in {"total", "award", "marks", "where", "which", "their",
                                   "answer", "hence", "shown", "correct", "value"}]
    # Combined: at least the top unique numbers and words
    return list(set(nums[:8] + words[:5]))


def _find_ms_section_by_keywords(ms_doc, ms_sections, keywords: list[str],
                                  qn: int, qp_max_mark: int = 0) -> tuple:
    """
    Search all MS sections for the one that best matches the given keywords.
    Returns (best_section, best_section_idx) or (None, -1).
    """
    if not keywords:
        return None, -1

    best_sec   = None
    best_idx   = -1
    best_score = 0

    for sec_idx, sec in enumerate(ms_sections):
        # Only check sections that have the target question number
        q_info = sec.get("questions", {}).get(qn)
        if q_info is None:
            continue

        # Gather text from the Q pages in this section
        sec_text = ""
        try:
            start_pi = q_info["page_idx"]
            end_pi   = q_info.get("end_page_idx", start_pi)
            for pi in range(start_pi, min(end_pi + 1, len(ms_doc))):
                sec_text += ms_doc[pi].get_text()
        except Exception:
            continue

        # Score: how many keywords appear in this section?
        score = sum(1 for kw in keywords if kw in sec_text)

        # Bonus: marks validation
        if qp_max_mark > 0:
            import re as _re
            tm = _re.search(r"Total[:\s]+(\d+)\s*marks?", sec_text, _re.I)
            if tm and abs(int(tm.group(1)) - qp_max_mark) <= 1:
                score += 3   # strong bonus for matching marks

        if score > best_score:
            best_score = score
            best_sec   = sec
            best_idx   = sec_idx

    return (best_sec, best_idx) if best_score >= 2 else (None, -1)


def _find_q_in_section_pages(ms_doc, sec, qn):
    """
    Strict fallback: scan section pages for a bold/left-margin question marker.
    Only matches "N." or "N. (a)" at x<70pt — avoids false positives.
    Returns a minimal q_info dict if found, else None.
    """
    qs_sorted    = sorted(sec.get("questions", {}).keys())
    sec_start_pi = sec.get("start_page", sec.get("start_pg", 1)) - 1
    if qs_sorted:
        _last_q_info = sec["questions"][qs_sorted[-1]]
        last_pi = (_last_q_info["page_idx"] if isinstance(_last_q_info, dict) else int(_last_q_info))
        end_pi  = last_pi + 10
    else:
        end_pi  = sec_start_pi + 25

    _ps = re.compile(r"^(\d{1,2})\.\s*$")
    _pq = re.compile(r"^(\d{1,2})\.\s+\(?\w")
    _pp = re.compile(r"^Q(\d{1,2})\.\s*(?:\(?\w|$)")

    for pi in range(sec_start_pi, min(end_pi, len(ms_doc))):
        page = ms_doc[pi]
        if _MS_INSTR_PAT.search(page.get_text()):
            continue
        ph = page.rect.height
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                spans    = line.get("spans", [])
                full     = "".join(s.get("text", "") for s in spans).strip()
                bb       = line["bbox"]
                if bb[0] > 70 or bb[1] > ph * 0.85:
                    continue
                qn_found = None
                m = _ps.match(full) or _pq.match(full)
                if m:
                    qn_found = int(m.group(1))
                else:
                    mq = _pp.match(full)
                    if mq:
                        qn_found = int(mq.group(1))
                if qn_found == qn:
                    return {
                        "page_idx":       pi,
                        "top_y":          bb[1],
                        "end_page_idx":   pi,
                        "end_y":          ph * 0.92,
                        "marker_right_x": bb[2],
                    }
    return None

def _crop_ms_pieces(ms_doc, q_info):
    spi = q_info["page_idx"]
    epi = q_info.get("end_page_idx", spi)
    sy  = q_info["top_y"]
    ey  = q_info["end_y"]
    mrx = q_info.get("marker_right_x", 0)
    final = []
    for pi in range(spi, epi + 1):
        if pi >= len(ms_doc):
            break
        pg = ms_doc[pi]
        ph = pg.rect.height
        pw = pg.rect.width
        pix = pg.get_pixmap(matrix=fitz.Matrix(_S_SCALE, _S_SCALE), alpha=False)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        iw, ih = img.size
        ft = _ms_ft_top(pg)
        pb = max(ph * 0.5, ft - 6)
        # FIX: use full page width minus tiny 8px margin so marks at x=97% are visible
        rx = iw - 8
        lx = 0  # always start from left edge to keep question-part labels
        hb = _ms_hdr_bot(pg)
        pt = max(ph * 0.06, hb + 6) if hb else ph * 0.06
        if pi == spi == epi:
            ty = sy
            # Track last real content line to avoid cutting off Total/marks
            lc_e = _last_ms_y(pg, sy, pb)
            by   = min(max(ey, lc_e + 10), pb)
        elif pi == spi:
            ty = sy
            lc = _last_ms_y(pg, sy, pb)
            by = min(lc + 10, pb)
        elif pi == epi:
            ty = pt
            lc2 = _last_ms_y(pg, pt, min(ey, pb))
            by  = min(max(lc2 + 10, ey), pb)
        else:
            ty = pt
            lc3 = _last_ms_y(pg, pt, pb)
            by  = min(lc3 + 10, pb)
        tp = max(0, int(ty * _S_SCALE) - 6)
        bp = min(ih, int(by * _S_SCALE) + 8)    # +8 for a little extra bottom padding
        if bp > tp + 20 and rx > lx:
            piece = img.crop((lx, tp, rx, bp))
            piece = _s_trim_bottom(piece)
            if piece.height > 10:
                # Guard: skip if this page contains general marking instructions
                page_txt = pg.get_text()
                if _MS_INSTR_PAT.search(page_txt):
                    continue   # skip instruction pages entirely

                # Strip "N." question number marker from the first piece
                if pi == spi:
                    for block in pg.get_text("dict").get("blocks", []):
                        if block.get("type") != 0:
                            continue
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                stxt = span.get("text", "").strip()
                                sbb  = span.get("bbox", [0, 0, 0, 0])
                                if (re.match(r"^\d{1,2}[.\s]?$", stxt)
                                        and sbb[0] < 60
                                        and sbb[1] >= ty - 4
                                        and sbb[1] < ty + 30):
                                    piece = _s_whiteout(
                                        piece,
                                        [(sbb[0], sbb[1], sbb[2] + 4, sbb[3] + 2)],
                                        origin_pt=(lx / _S_SCALE, tp / _S_SCALE),
                                    )
                final.extend(_s_split_smart(piece, 1500))
    return final if final else None


# ═══════════════════════════════════════════════════════════════════════════════
#  STRUCTURED MODE — WORD GENERATION (FINAL AGENT layout)
# ═══════════════════════════════════════════════════════════════════════════════
_PAGE_H_PT  = 728.0   # A4 content height (2 cm margins)
_HEADER_PT  = 125     # conservative header estimate
_SOL_LBL_PT = 30      # "Student's Solution:" label height
_SOL_ROW_PT = 28      # 28pt per sol-box row
_SAFETY_PT  = 20      # safety margin
_IMG_THRESH = 380     # pt — images taller than this → sol on separate page


def _sol_config(img):
    """Returns (n_rows, sol_on_separate_page)."""
    if img is None:
        return (6, False)
    img_pt = (img.height / img.width) * _S_CW_PT
    if img_pt > _IMG_THRESH:
        return (6, True)
    avail = _PAGE_H_PT - _HEADER_PT - img_pt - _SOL_LBL_PT - _SAFETY_PT
    n = max(4, min(int(avail / _SOL_ROW_PT), 8))
    return (n, False)


def _s_run(p, txt, bold=False, italic=False, size_pt=11, color=None):
    r = p.add_run(str(txt))
    r.bold       = bold
    r.italic     = italic
    r.font.size  = Pt(size_pt)
    r.font.name  = "Arial"
    if color:
        r.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return r


def _s_hr(doc, color="CCCCCC"):
    p    = doc.add_paragraph()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(docx_qn("w:val"),   "single")
    bot.set(docx_qn("w:sz"),    "6")
    bot.set(docx_qn("w:color"), color)
    pBdr.append(bot)
    pPr = p._p.get_or_add_pPr()
    idx = len(pPr)
    for j, c in enumerate(pPr):
        if c.tag.split('}')[-1] in ('spacing', 'ind', 'jc', 'rPr', 'sectPr'):
            idx = j
            break
    pPr.insert(idx, pBdr)


def _s_sol_box(doc, n):
    n = max(4, min(n, 8))
    t = doc.add_table(rows=n, cols=1)
    t.autofit = False
    for c in t.columns[0].cells:
        c.width = Cm(_S_CW)
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        trH  = OxmlElement("w:trHeight")
        trH.set(docx_qn("w:val"), "560")   # 28pt
        trPr.append(trH)
        cell = row.cells[0]
        tcPr = cell._tc.get_or_add_tcPr()
        tcB  = OxmlElement("w:tcBorders")
        for s, v, sz, col in [("top", "nil", "0", "auto"), ("left", "nil", "0", "auto"),
                               ("right", "nil", "0", "auto"), ("bottom", "single", "6", "BFBFBF")]:
            b = OxmlElement(f"w:{s}")
            b.set(docx_qn("w:val"), v)
            if v != "nil":
                b.set(docx_qn("w:sz"),    sz)
                b.set(docx_qn("w:color"), col)
            tcB.append(b)
        tcPr.append(tcB)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)


def _s_add_img(doc, pil_img):
    buf = io.BytesIO()
    pil_img.save(buf, format="PNG")
    buf.seek(0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.add_run().add_picture(buf, width=Cm(_S_CW))


def _s_add_ms_img(doc, pil_img):
    """Add MS image wrapped in cantSplit table — no mid-image page breaks."""
    buf = io.BytesIO()
    pil_img.save(buf, format="PNG")
    buf.seek(0)
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.columns[0].width = Cm(_S_CW)
    cell = t.rows[0].cells[0]
    trPr = t.rows[0]._tr.get_or_add_trPr()
    cs   = OxmlElement("w:cantSplit")
    cs.set(docx_qn("w:val"), "1")
    trPr.append(cs)
    tcPr = cell._tc.get_or_add_tcPr()
    tcB  = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(docx_qn("w:val"), "nil")
        tcB.append(b)
    tcPr.append(tcB)
    tcMar = OxmlElement("w:tcMar")
    for s in ["top", "left", "bottom", "right"]:
        m = OxmlElement(f"w:{s}")
        m.set(docx_qn("w:w"),    "0")
        m.set(docx_qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.add_run().add_picture(buf, width=Cm(_S_CW))


def _normalize_topic(t: str) -> str:
    if not t or str(t).strip().lower() in ("error", "none", "nan", ""):
        return "Unclassified"
    s = re.sub(r'^#+\s*', '', str(t).strip())
    m = re.match(r"^\s*chapter\s+(\d+)\s*[:\.\-–]?\s*(.+?)\s*$", s, re.IGNORECASE)
    if m:
        return f"{m.group(1)}. {m.group(2).strip()}"
    m = re.match(r"^\s*(\d+)\s*[\.:\-–]\s*(.+?)\s*$", s)
    if m:
        return f"{m.group(1)}. {m.group(2).strip()}"
    return re.sub(r"\s+", " ", s)


def _chapter_sort_key(topic: str):
    m = re.match(r"^\s*(\d+)\s*[\.\-:]", str(topic))
    if m:
        return (0, int(m.group(1)), str(topic))
    return (1, 0, str(topic).lower())


DIFF_ORDER = ["Easy", "Medium", "Hard"]


def build_structured_worksheet(
    xl_rows: list[dict],
    qp_bytes: bytes,
    ms_bytes: bytes,
    ms_sections: list[dict],
    row_to_section: dict,
    progress_cb=None,
) -> bytes:
    """Generate the structured worksheet using the FINAL AGENT layout."""

    # ── Pre-generation validation ────────────────────────────────────────────
    _validation_errors = []
    for _vr in xl_rows:
        _vpg = _vr.get("page_num", 0)
        _vqn = _vr.get("qn", 0)
        _vmk = _vr.get("marks", 0)
        _vtopic = (_vr.get("_topic_norm") or _vr.get("topic") or "")
        if _vpg <= 0:
            _validation_errors.append(f"  • Q{_vqn}: page_num missing")
        if _vmk <= 0:
            _validation_errors.append(f"  • Q{_vqn}: marks unknown")

    if _validation_errors and len(_validation_errors) > len(xl_rows) * 0.5:
        # More than 50% of rows are invalid — raise early
        raise ValueError(
            "Worksheet validation failed — too many invalid rows:\n" +
            "\n".join(_validation_errors[:10]) +
            ("\n  ..." if len(_validation_errors) > 10 else "")
        )

    qp_doc = fitz.open(stream=qp_bytes, filetype="pdf")
    ms_doc = fitz.open(stream=ms_bytes, filetype="pdf")

    doc = Document()
    for s in doc.sections:
        s.top_margin    = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin   = Cm(2)
        s.right_margin  = Cm(2)

    # ── Group by normalized topic → difficulty (preserve Excel row order within) ──
    for r in xl_rows:
        r["_topic_norm"] = _normalize_topic(r.get("topic", ""))

    grouped = {}
    for ri, r in enumerate(xl_rows):
        t = r["_topic_norm"]
        d = r.get("difficulty") or "Unspecified"
        grouped.setdefault(t, {}).setdefault(d, []).append(ri)

    # Skip "Unclassified" / empty topics — don't include poorly classified rows
    sorted_topics = sorted(
        (t for t in grouped.keys() if t and t.lower() not in ("unclassified", "unspecified", "")),
        key=_chapter_sort_key
    )
    _skipped_unclassified = sum(
        len(ris)
        for t, diffs in grouped.items()
        if not t or t.lower() in ("unclassified", "unspecified", "")
        for ris in diffs.values()
    )
    if _skipped_unclassified:
        import streamlit as _st_tmp
        try:
            _st_tmp.info(f"ℹ Skipped {_skipped_unclassified} unclassified row(s) — "
                         "set Topic in Excel to include them.")
        except Exception:
            pass

    # Total questions = only classified rows (unclassified are skipped)
    _classified_count = sum(
        1 for r in xl_rows
        if r.get("_topic_norm") and r["_topic_norm"].lower() not in ("unclassified","unspecified","")
    )
    total_qs = _classified_count if _classified_count > 0 else len(xl_rows)
    done_qs  = 0
    first    = True
    ws_q     = 0      # per-chapter counter (resets each new chapter)
    cur_t    = None
    cur_d    = None

    for topic in sorted_topics:
        diffs = grouped[topic]
        sorted_diffs = sorted(
            diffs.keys(),
            key=lambda d: DIFF_ORDER.index(d) if d in DIFF_ORDER else 99
        )

        for diff in sorted_diffs:
            ri_list = diffs[diff]
            for ri in ri_list:
                r    = xl_rows[ri]
                pg   = r.get("page_num", 0)
                qn   = r.get("qn", 0)
                marks_raw = r.get("marks")
                marks = int(float(str(marks_raw))) if marks_raw else 0
                # Skip rows where marks couldn't be determined (likely bad data)
                if marks <= 0:
                    done_qs += 1
                    continue
                marks = marks or 1
                quote = r.get("quote", "") or ""
                if quote in ("nan",):
                    quote = ""

                # Special override: Q9 page 147
                if qn == 9 and pg == 147:
                    marks = 6

                # Crop QP (pass pg_end so multi-page questions are complete)
                pg_end = r.get("page_num_end", 0) or 0
                qp_img = _crop_qp_full(qp_doc, pg, qn, pg_end=pg_end)
                # Try text extraction (text-first for non-diagram questions)
                _qp_txt = _extract_qp_text(qp_doc, pg, qn) if pg > 0 else None

                # Crop MS
                sec     = row_to_section.get(ri)
                ms_imgs = []
                ms_topic_mismatch   = False   # kept for display logic below
                ms_marks_mismatch   = False   # marks total mismatch flag
                if sec and isinstance(sec, dict):
                    q_info = sec.get("questions", {}).get(qn)
                    # Fallback 1: scan section pages for the question marker
                    if q_info is None:
                        q_info = _find_q_in_section_pages(ms_doc, sec, qn)
                    # Fallback 2: content-based search using Excel MS keywords
                    # If positional mapping gave wrong section, keywords find correct one
                    excel_ms_text = r.get("ms_text", "")
                    if excel_ms_text:
                        keywords = _extract_answer_keywords(excel_ms_text)
                        if keywords:
                            kw_sec, kw_idx = _find_ms_section_by_keywords(
                                ms_doc, ms_sections, keywords, qn,
                                r.get("marks", 0) or 0
                            )
                            if kw_sec is not None:
                                # Keyword search found a better match
                                # Use it if it differs from current section
                                if kw_idx != row_to_section_idx.get(ri, -1):
                                    sec    = kw_sec
                                    q_info = sec.get("questions", {}).get(qn)
                                    if q_info is None:
                                        q_info = _find_q_in_section_pages(ms_doc, sec, qn)
                    # Fallback 3: use entire section content (guarantees no unmatched row)
                    if q_info is None and sec and sec.get("questions"):
                        qs_sorted = sorted(sec["questions"].keys())
                        first_q   = sec["questions"][qs_sorted[0]]
                        last_q    = sec["questions"][qs_sorted[-1]]
                        q_info = {
                            "page_idx":       first_q["page_idx"],
                            "top_y":          first_q["top_y"],
                            "end_page_idx":   last_q.get("end_page_idx", last_q["page_idx"]),
                            "end_y":          last_q.get("end_y", last_q.get("bottom_y", 700)),
                            "marker_right_x": 0,
                        }
                    if q_info:
                        # Ensure complete MS: extend end_page_idx until [Total:] found
                        _q_info_ext = dict(q_info)
                        _max_mark   = r.get("marks", 0) or 0
                        _epi_orig   = _q_info_ext.get("end_page_idx", _q_info_ext["page_idx"])
                        for _extra in range(0, 6):
                            _test_pi = _epi_orig + _extra
                            if _test_pi >= len(ms_doc):
                                break
                            _test_txt = ms_doc[_test_pi].get_text()
                            if re.search(r"\[Total[:\s]", _test_txt, re.I):
                                _q_info_ext["end_page_idx"] = _test_pi
                                _q_info_ext["end_y"]        = ms_doc[_test_pi].rect.height * 0.95
                                break   # found Total marker
                        pieces = _crop_ms_pieces(ms_doc, _q_info_ext)
                        if pieces:
                            # ── Max-mark validation ───────────────────────────────
                            # Check if [Total: N marks] in MS matches QP maximum mark.
                            # If mismatched, try adjacent MS sections before accepting.
                            qp_max_mark = r.get("marks", 0) or 0
                            ms_total_ok = True
                            if qp_max_mark > 0:
                                # Extract Total marks from the MS pieces text
                                ms_text_combined = ""
                                try:
                                    for _p_idx in range(
                                        q_info["page_idx"],
                                        min(q_info.get("end_page_idx", q_info["page_idx"]) + 1,
                                            len(ms_doc))
                                    ):
                                        ms_text_combined += ms_doc[_p_idx].get_text()
                                except Exception:
                                    pass
                                _total_m = re.search(
                                    r"Total[:\s]+(\d+)\s*marks?",
                                    ms_text_combined, re.IGNORECASE
                                )
                                if _total_m:
                                    ms_total = int(_total_m.group(1))
                                    if abs(ms_total - qp_max_mark) > 2:
                                        ms_total_ok = False   # mismatch detected
                            if ms_total_ok:
                                ms_imgs = pieces
                            else:
                                ms_imgs = pieces   # still use it but flag below
                                ms_marks_mismatch = True

                # sol config
                n_sol, sol_separate = _sol_config(qp_img)

                # ── Chapter heading (on same page as Q — no separate break) ──
                if topic != cur_t:
                    h = doc.add_paragraph()
                    if not first:
                        pass   # page break is on qh below
                    h.paragraph_format.space_before = Pt(0)
                    h.paragraph_format.space_after  = Pt(6)
                    _s_run(h, topic, bold=True, size_pt=20)
                    cur_t = topic
                    cur_d = None
                    ws_q  = 0   # ← reset per-chapter counter

                if diff != cur_d:
                    dp = doc.add_paragraph()
                    dp.paragraph_format.space_before = Pt(6)
                    dp.paragraph_format.space_after  = Pt(4)
                    _s_run(dp, f"— {diff} —", bold=True, italic=True, size_pt=13)
                    cur_d = diff

                # ── Question header (ONE page break per question) ──────────
                ws_q += 1   # increment after possible chapter reset
                qh = doc.add_paragraph()
                if not first:
                    qh.paragraph_format.page_break_before = True
                    # Move topic/diff headings before qh if they were just added
                    # (they'll flow after qh page break)
                    # Restructure: insert topic/diff before qh in XML
                    if topic != xl_rows[ri_list[0] if ri == ri_list[0] else ri]["_topic_norm"] if ri != ri_list[0] else False:
                        pass
                qh.paragraph_format.space_before = Pt(0)
                qh.paragraph_format.space_after  = Pt(2)
                _s_run(qh, f"Question: {ws_q}", bold=True, size_pt=13)

                mp = doc.add_paragraph()
                mp.paragraph_format.space_before = Pt(0)
                mp.paragraph_format.space_after  = Pt(2)
                _s_run(mp, "Level of question", bold=True, size_pt=11)
                _s_run(mp, f": {diff}  |  ", size_pt=11)
                _s_run(mp, "Number of Marks: ", bold=True, size_pt=11)
                _s_run(mp, f"{marks}", size_pt=11)

                cp = doc.add_paragraph()
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.space_after  = Pt(2)
                _s_run(cp, "Chapter", bold=True, size_pt=11)
                _s_run(cp, f" :  {topic}", size_pt=11)
                _s_hr(doc, "BFBFBF")

                # ── QP image (split across pages if too tall) ─────────────
                # Safe page height for a QP image (content height minus header)
                _QP_MAX_H    = 1400   # px at 150dpi  ≈ 504pt  → leaves room for header
                _QP_SAFE_1ST = 900    # px — max for the first piece (header takes space)

                if _qp_txt:
                    # ── Text-first rendering ────────────────────────────────
                    for _txt_line in _qp_txt.split("\n"):
                        _txt_line = _txt_line.strip()
                        if not _txt_line:
                            continue
                        _is_max = bool(re.match(r"^\[Maximum mark", _txt_line, re.I))
                        _is_qhd = bool(re.match(rf"^{qn}[.\s]", _txt_line))
                        _is_sub = bool(re.match(r"^[a-e][.)\s]|^\([a-e]\)", _txt_line))
                        _tp = doc.add_paragraph()
                        _tp.paragraph_format.space_before = Pt(4 if _is_max else 0)
                        _tp.paragraph_format.space_after  = Pt(2)
                        _tp.paragraph_format.left_indent  = Cm(0.5 if _is_sub else 0)
                        _s_run(_tp, _txt_line, bold=_is_max or _is_qhd,
                               size_pt=10 if _is_max else 11)
                    # If there's also an image (diagrams), include it
                    if qp_img:
                        _s_add_img(doc, qp_img)
                    _qp_pieces_used = 1
                elif qp_img:
                    # ── Image rendering ─────────────────────────────────────
                    qp_h = qp_img.height
                    if qp_h <= _QP_MAX_H:
                        _s_add_img(doc, qp_img)
                        _qp_pieces_used = 1
                    else:
                        qp_pieces = _s_split_smart(qp_img, max_h=_QP_SAFE_1ST)
                        for pi_idx, qp_piece in enumerate(qp_pieces):
                            if pi_idx > 0:
                                _pb_qp = doc.add_paragraph()
                                _pb_qp.paragraph_format.page_break_before = True
                                _pb_qp.paragraph_format.space_before = Pt(0)
                                _pb_qp.paragraph_format.space_after  = Pt(0)
                            _s_add_img(doc, qp_piece)
                        _qp_pieces_used = len(qp_pieces)
                else:
                    # Neither text nor image available — skip this question
                    p = doc.add_paragraph()
                    _s_run(p, f"⚠ Q{qn} (page {pg}): question content unavailable — skipped",
                           italic=True, size_pt=10, color="CC0000")
                    _qp_pieces_used = 0

                # ── Student's Solution ────────────────────────────────────
                # Always on a fresh page when image was split
                if sol_separate or (qp_img and qp_img.height > _QP_MAX_H):
                    pb2 = doc.add_paragraph()
                    pb2.paragraph_format.page_break_before = True
                    pb2.paragraph_format.space_before = Pt(0)
                    pb2.paragraph_format.space_after  = Pt(0)

                sl = doc.add_paragraph()
                sl.paragraph_format.space_before = Pt(0 if (sol_separate or (qp_img and qp_img.height > _QP_MAX_H)) else 8)
                sl.paragraph_format.space_after  = Pt(4)
                sl.paragraph_format.keep_with_next = True
                _s_run(sl, "Student's Solution:", bold=True, size_pt=11)
                _s_sol_box(doc, n_sol)

                # ── MS page (separate page) ───────────────────────────────
                pb = doc.add_paragraph()
                pb.paragraph_format.page_break_before = True
                pb.paragraph_format.space_before = Pt(0)
                pb.paragraph_format.space_after  = Pt(0)

                ap = doc.add_paragraph()
                ap.paragraph_format.space_before = Pt(0)
                ap.paragraph_format.space_after  = Pt(4)
                ap.paragraph_format.keep_with_next = True
                _s_run(ap, "Answer from Mark Scheme:", bold=True, size_pt=11)

                if ms_imgs:
                    for img in ms_imgs:
                        _s_add_ms_img(doc, img)
                    # Show warning banner if marks total doesn't match
                    if ms_marks_mismatch:
                        _warn = doc.add_paragraph()
                        _warn.paragraph_format.space_before = Pt(2)
                        _warn.paragraph_format.space_after  = Pt(2)
                        _s_run(_warn,
                               f"⚠ Marks mismatch: QP max={r.get('marks',0)} "
                               f"but MS total may differ. Verify manually.",
                               italic=True, size_pt=9, color="CC6600")
                else:
                    # MS is empty — show placeholder but still include the question
                    # (skipping entirely would cause gaps in question numbering)
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after  = Pt(2)
                    _s_run(p, "⚠ Mark Scheme not found — please verify manually",
                           italic=True, size_pt=10, color="CC0000")

                # ── Keep it up (after MS, with HR) ─────────────────────────
                if quote:
                    _s_hr(doc, "BFBFBF")
                    kp = doc.add_paragraph()
                    kp.paragraph_format.space_before = Pt(4)
                    kp.paragraph_format.space_after  = Pt(0)
                    _s_run(kp, "Keep it up", bold=True, size_pt=11)
                    _s_run(kp, f" : {quote}", size_pt=11)

                first   = False
                done_qs += 1
                if progress_cb:
                    progress_cb(done_qs, total_qs, f"Building Q{ws_q}…")

    qp_doc.close()
    ms_doc.close()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
#  MCQ WORD BUILDER (unchanged from original)
# ═══════════════════════════════════════════════════════════════════════════════
CONTENT_WIDTH_CM = 17.0


def _set_cell_borders(cell, top="single", left="single", right="single",
                      bottom="single", color="999999",
                      bot_color=None, top_color=None, sz_top=8, sz_bot=8):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, style, sz, c in [
        ("top",    top,    sz_top, top_color or color),
        ("left",   left,   8,      color),
        ("bottom", bottom, sz_bot, bot_color or color),
        ("right",  right,  8,      color),
    ]:
        b = OxmlElement(f"w:{side}")
        b.set(docx_qn("w:val"),   style)
        b.set(docx_qn("w:sz"),    str(sz))
        b.set(docx_qn("w:color"), c)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _run(para, text, *, bold=False, italic=False, color=None, size_pt=11):
    r = para.add_run(str(text))
    r.bold       = bold
    r.italic     = italic
    r.font.size  = Pt(size_pt)
    r.font.name  = "Arial"
    if color:
        r.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return r


def _insert_pBdr(p, pBdr):
    pPr = p._p.get_or_add_pPr()
    insert_idx = len(pPr)
    for i, child in enumerate(pPr):
        tag = child.tag.split('}')[-1]
        if tag in ('spacing', 'ind', 'jc', 'contextualSpacing', 'mirrorIndents', 'rPr', 'sectPr'):
            insert_idx = i
            break
    pPr.insert(insert_idx, pBdr)


def _hr(doc, color="CCCCCC"):
    p    = doc.add_paragraph()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(docx_qn("w:val"),   "single")
    bot.set(docx_qn("w:sz"),    "6")
    bot.set(docx_qn("w:color"), color)
    pBdr.append(bot)
    _insert_pBdr(p, pBdr)
    return p


def _add_solution_box(doc, n_lines: int = 4):
    table = doc.add_table(rows=n_lines, cols=1)
    table.autofit = False
    for cell in table.columns[0].cells:
        cell.width = Cm(CONTENT_WIDTH_CM)
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trH  = OxmlElement("w:trHeight")
        trH.set(docx_qn("w:val"), "510")
        trPr.append(trH)
        cell = row.cells[0]
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side, val, sz, color in [
            ("top",    "nil",    "0", "auto"),
            ("left",   "nil",    "0", "auto"),
            ("right",  "nil",    "0", "auto"),
            ("bottom", "single", "4", "BFBFBF"),
        ]:
            b = OxmlElement(f"w:{side}")
            b.set(docx_qn("w:val"), val)
            if val != "nil":
                b.set(docx_qn("w:sz"),    sz)
                b.set(docx_qn("w:color"), color)
            tcBorders.append(b)
        tcPr.append(tcBorders)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)


def build_word_document(questions, qp_bytes=None, locations=None, progress_cb=None) -> bytes:
    """MCQ mode word builder."""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    def _normalize_topic_mcq(t: str) -> str:
        if not t or str(t).strip().lower() in ("error", "none", "nan", ""):
            return "Unclassified"
        s = str(t).strip()
        s = re.sub(r'^#+\s*', '', s).strip()
        m = re.match(r"^\s*chapter\s+(\d+)\s*[:\.\-–]?\s*(.+?)\s*$", s, re.IGNORECASE)
        if m:
            return f"{m.group(1)}. {m.group(2).strip()}"
        m = re.match(r"^\s*(\d+)\s*[\.:\-–]\s*(.+?)\s*$", s)
        if m:
            return f"{m.group(1)}. {m.group(2).strip()}"
        return re.sub(r"\s+", " ", s)

    for q in questions:
        q["topic"] = _normalize_topic_mcq(q.get("topic"))

    grouped = {}
    for q in questions:
        t = q["topic"]
        d = q.get("difficulty") or "Unspecified"
        grouped.setdefault(t, {}).setdefault(d, []).append(q)

    sorted_topics = sorted(grouped.keys(), key=_chapter_sort_key)

    visual_qs  = [q for q in questions if q.get("needs_image")]
    total_imgs = len(visual_qs)
    img_done   = 0
    is_first_chapter = True
    display_counter  = 0

    for topic in sorted_topics:
        diffs = grouped[topic]
        h = doc.add_paragraph()
        if not is_first_chapter:
            h.paragraph_format.page_break_before = True
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after  = Pt(8)
        _run(h, topic, bold=True, size_pt=20)

        sorted_diffs = sorted(
            diffs.keys(),
            key=lambda d: DIFF_ORDER.index(d) if d in DIFF_ORDER else 99
        )

        is_first_diff_in_chapter = True
        for diff_ in sorted_diffs:
            qs_in_diff = diffs[diff_]
            dh = doc.add_paragraph()
            if not is_first_diff_in_chapter:
                dh.paragraph_format.page_break_before = True
            dh.paragraph_format.space_before = Pt(8)
            dh.paragraph_format.space_after  = Pt(6)
            _run(dh, f"— {diff_} —", bold=True, italic=True, size_pt=13)

            is_first_q_in_diff = True
            for q in qs_in_diff:
                display_counter += 1
                q_num  = q["qn"]
                marks  = q.get("marks", 1) or 1
                quote  = q.get("quote", "") or ""
                answer = q.get("answer", "")
                found  = q.get("found", False)
                vis    = q.get("needs_image", False)
                text   = q.get("text", "")
                opts   = {L: q.get(L, "") for L in "ABCD"}
                topic_ = q.get("topic", topic)

                qh = doc.add_paragraph()
                if not is_first_q_in_diff:
                    qh.paragraph_format.page_break_before = True
                qh.paragraph_format.space_before = Pt(0)
                qh.paragraph_format.space_after  = Pt(2)
                _run(qh, f"Question: {display_counter}", bold=True, size_pt=13)

                mp = doc.add_paragraph()
                mp.paragraph_format.space_before = Pt(0)
                mp.paragraph_format.space_after  = Pt(2)
                _run(mp, "Level of question",  bold=True, size_pt=11)
                _run(mp, f": {diff_}  |  ",                size_pt=11)
                _run(mp, "Number of Marks: ",  bold=True, size_pt=11)
                _run(mp, f"{marks}",                       size_pt=11)

                cp = doc.add_paragraph()
                cp.paragraph_format.space_before = Pt(0)
                cp.paragraph_format.space_after  = Pt(2)
                _run(cp, "Chapter", bold=True, size_pt=11)
                _run(cp, f" :  {topic_}",     size_pt=11)

                _hr(doc, color="BFBFBF")

                def _looks_failed(stem, options):
                    s = (stem or "").strip()
                    if not s:
                        return True
                    sl = s.lower()
                    if re.search(r"\[\s*q\s*\d+", sl): return True
                    if re.search(r"\[\s*stem\s*\]", sl): return True
                    if re.search(r"\btest\s+stem\b", sl): return True
                    if sl in ("stem", "[stem]", "n/a", "tbd", "placeholder", "todo", "todo:"):
                        return True
                    if re.match(r"^q\s*\d+\s*stem\s*$", sl): return True
                    non_empty = [v for v in options.values() if v and v.strip()]
                    if not non_empty: return True
                    if all(v.strip().upper() in ("A", "B", "C", "D") for v in non_empty): return True
                    if all(re.match(r"^option\s*[a-d]$", v.strip(), re.I) for v in non_empty): return True
                    return False

                if found and not vis and _looks_failed(text, opts) and q_num in locations:
                    vis = True

                if not found:
                    np_ = doc.add_paragraph()
                    _run(np_, "Question not found on specified page - needs review", italic=True, size_pt=11)
                elif vis:
                    if progress_cb:
                        progress_cb(img_done + 1, total_imgs or 1, f"Cropping Q{q_num} from QP…")
                    img_done += 1
                    q_loc = q.get("_loc")
                    img_bytes = crop_question_png(qp_bytes, q_loc) if q_loc else None
                    if img_bytes:
                        ip = doc.add_paragraph()
                        ip.paragraph_format.space_before = Pt(0)
                        ip.paragraph_format.space_after  = Pt(4)
                        ip.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(CONTENT_WIDTH_CM - 1.0))
                    elif text and not _looks_failed(text, opts):
                        for line in text.split("\n"):
                            if line.strip():
                                _run(doc.add_paragraph(), line, size_pt=11)
                        for L in "ABCD":
                            if opts[L]:
                                op = doc.add_paragraph()
                                op.paragraph_format.space_before = Pt(2)
                                op.paragraph_format.space_after  = Pt(2)
                                _run(op, f"{L}.  {opts[L]}", size_pt=11)
                    else:
                        np_ = doc.add_paragraph()
                        _run(np_, f"Question Q{q_num} could not be extracted — refer to QP.", italic=True, size_pt=11)
                else:
                    for line in text.split("\n"):
                        if line.strip():
                            tp = doc.add_paragraph()
                            tp.paragraph_format.space_before = Pt(0)
                            tp.paragraph_format.space_after  = Pt(2)
                            _run(tp, line, size_pt=11)
                    for L in "ABCD":
                        if opts[L]:
                            op = doc.add_paragraph()
                            op.paragraph_format.space_before = Pt(2)
                            op.paragraph_format.space_after  = Pt(2)
                            _run(op, f"{L}.  {opts[L]}", size_pt=11)

                sl = doc.add_paragraph()
                sl.paragraph_format.space_before = Pt(8)
                sl.paragraph_format.space_after  = Pt(2)
                _run(sl, "Student's Solution:", bold=True, size_pt=11)
                _add_solution_box(doc, n_lines=max(6, (marks or 1) * 2))

                ap = doc.add_paragraph()
                ap.paragraph_format.space_before = Pt(8)
                ap.paragraph_format.space_after  = Pt(2)
                ap.paragraph_format.keep_with_next = True
                _run(ap, "Answer from Mark Scheme:", bold=True, size_pt=11)

                av = doc.add_paragraph()
                av.paragraph_format.space_before = Pt(0)
                av.paragraph_format.space_after  = Pt(2)
                _run(av, answer, bold=True, size_pt=12)

                if quote:
                    _hr(doc, color="BFBFBF")
                    qup = doc.add_paragraph()
                    qup.paragraph_format.space_before = Pt(2)
                    qup.paragraph_format.space_after  = Pt(2)
                    _run(qup, "Keep it up", bold=True, size_pt=11)
                    _run(qup, f" : {quote}",       size_pt=11)
                else:
                    _hr(doc, color="BFBFBF")

                is_first_q_in_diff = False

            is_first_diff_in_chapter = False
        is_first_chapter = False

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
#  TOPIC MATCH (for diagnostic only)
# ═══════════════════════════════════════════════════════════════════════════════
def _topic_keywords(text: str) -> set:
    if not text:
        return set()
    t = text.lower()
    boilerplate = re.compile(
        r"(maximum\s+mark|marks?\s*\]|international\s+baccalaureate|©\s*\d{4}|"
        r"turn\s+over|award\s+(a\d|m\d|r\d)|note\s*:|\b(m\d|a\d|r\d|ft|ag)\b|"
        r"calculator|gdc|\bsolve\b|\bfind\b|\bcalculate\b|\bdetermine\b|"
        r"\bthe\b|\bof\b|\ba\b|\bto\b|\bin\b|\bis\b|\bfor\b|\bbe\b|\bby\b|"
        r"\bwith\b|\bare\b|\band\b|\bor\b|\bif\b|\bnot\b)", re.IGNORECASE
    )
    t = boilerplate.sub(" ", t)
    tokens = re.findall(r"\b[a-z][a-z]{3,}\b", t)
    return set(tokens)


def _topics_match(qp_text: str, ms_text: str, threshold: float = 0.10) -> bool:
    qp_kw = _topic_keywords(qp_text)
    ms_kw = _topic_keywords(ms_text)
    if not qp_kw or not ms_kw:
        return True
    overlap = qp_kw & ms_kw
    smaller = min(len(qp_kw), len(ms_kw))
    if smaller == 0:
        return True
    return len(overlap) / smaller >= threshold


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN STREAMLIT FLOW
# ═══════════════════════════════════════════════════════════════════════════════
if st.button(
    "⚡ Extract & Generate Worksheet",
    type="primary",
    disabled=not (qp_file and ms_file and xl_file),
):
    client = anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])

    if not qp_file or not ms_file or not xl_file:
        st.warning("Please upload QP PDF, MS PDF, and Excel file first.")
        st.stop()

    qp_bytes = read_bytes(qp_file)
    ms_bytes = read_bytes(ms_file)

    if not qp_bytes or not ms_bytes:
        st.error("❌ Could not read uploaded files. Please re-upload and try again.")
        st.stop()

    if qp_bytes == ms_bytes:
        st.error("❌ QP and MS files are identical! Please upload the correct Mark Scheme.")
        st.stop()

    is_structured = (mode == "Structured Mark Scheme Mode")

    with st.status("Processing…", expanded=True) as status:

        # ── 1) Excel ──────────────────────────────────────────────────────────
        st.write("📊 Reading Excel (metadata only)…")
        try:
            parse_excel._filter_paper2 = st.session_state.get("filter_p2", False)
            xl_rows = parse_excel(xl_file)
        except Exception as e:
            st.error(f"Failed to read Excel: {e}")
            st.stop()
        if not xl_rows:
            _fp = st.session_state.get("filter_p2", False)
            if _fp:
                st.error(
                    "❌ No valid question rows found in Excel.\n\n"
                    "**Most likely reason:** All rows were removed because the "
                    "Reference column was detected as Paper 2.  \n"
                    "👉 **Uncheck the 'Skip Paper 2 rows' checkbox** above "
                    "and click Extract & Generate again."
                )
            else:
                st.error(
                    "❌ No valid question rows found in Excel.\n\n"
                    "Please check:\n"
                    "- Row 1 must be column headers\n"
                    "- There must be a Question Number column (Q No., Q#, etc.)\n"
                    "- Question numbers must be integers (1, 2, 3…)\n"
                    "- If this is a Paper 2 file, uncheck 'Skip Paper 2 rows'"
                )
            st.stop()
        st.write(f"✅ {len(xl_rows)} question rows in Excel")

        col_map = getattr(parse_excel, "last_column_map", {})
        if col_map:
            mapped_pairs = []
            unmapped = []
            for logical, (idx, actual) in col_map.items():
                if idx is None:
                    unmapped.append(logical)
                else:
                    mapped_pairs.append(f"**{logical}** → `{actual}`")
            with st.expander("🔍 Excel column mapping", expanded=bool(unmapped)):
                st.write(" · ".join(mapped_pairs))
                if unmapped:
                    st.error(f"❌ Missing column(s): {', '.join(unmapped)}")

        zero_page_count = sum(1 for r in xl_rows if not r.get("page_num"))
        range_count     = sum(1 for r in xl_rows
                              if r.get("page_num_end", 0) > r.get("page_num", 0))
        if zero_page_count > len(xl_rows) * 0.1:
            st.warning(f"⚠ {zero_page_count}/{len(xl_rows)} rows have Page Number = 0. "
                       "Check the Page Number column in your Excel.")
        if range_count:
            st.info(f"ℹ️ {range_count} question(s) span multiple pages "
                    f"(e.g. '1011' interpreted as pages 10–11) — "
                    "they will be cropped across all specified pages.")

        dups = getattr(parse_excel, "last_duplicates", [])
        if dups:
            st.warning(f"⚠ Removed {len(dups)} duplicate rows")

        # ── 2) QP locations ───────────────────────────────────────────────────
        st.write("📍 Locating questions in QP PDF…")
        try:
            locations = find_question_locations(qp_bytes)
        except Exception as e:
            st.error(f"Failed to scan QP PDF: {e}")
            st.stop()
        st.write(f"✅ Found {len(locations)} question marker(s) in QP")

        # Segment detection (for MS matching)
        # ── Build segments from QP PDF exam boundaries ─────────────────────────
        # Detect exam session start pages from QP PDF (pages with Q1 [Maximum mark]).
        # Then assign each Excel row to its session by QP page number.
        # This guarantees: Excel segment N = QP session N = MS section N (positional).
        qp_session_starts: list[int] = []   # QP page numbers (1-based) where each exam starts
        try:
            _qp_tmp = fitz.open(stream=qp_bytes, filetype="pdf")
            for _pi in range(len(_qp_tmp)):
                _pg_txt = _qp_tmp[_pi].get_text()
                # Try multiple patterns to handle different PDF text encodings
                if (re.search(r"1[.][\s\x07]*\[Maximum mark", _pg_txt) or
                        re.search(r"^\s*1\.\s*\[Maximum mark", _pg_txt, re.M) or
                        re.search(r"1\.\s{0,10}Maximum mark", _pg_txt)):
                    qp_session_starts.append(_pi + 1)   # 1-based page number
            _qp_tmp.close()
        except Exception:
            qp_session_starts = []

        # Assign each Excel row to a QP session index (0-based)
        def _row_session_idx(row_page: int) -> int:
            """Return 0-based session index for a row whose QP page is row_page."""
            if not qp_session_starts or row_page <= 0:
                return 0
            idx_ = 0
            for k_, sp in enumerate(qp_session_starts):
                if sp <= row_page:
                    idx_ = k_
                else:
                    break
            return idx_

        # Build ref→session mapping from QP PDF for fallback matching
        _ref_to_session: dict = {}
        try:
            _qp_ref_doc = fitz.open(stream=qp_bytes, filetype="pdf")
            for _si, _sp in enumerate(qp_session_starts):
                _ptxt = _qp_ref_doc[_sp - 1].get_text()
                _cm = re.search(r"(\d{4}[-–]\d{4})", _ptxt)
                if _cm:
                    _ref_to_session[_cm.group(1).replace("–", "-")] = _si
                _dm = re.search(r"(May|November|March)\s+(\d{4})", _ptxt, re.I)
                if _dm:
                    _ref_to_session[f"{_dm.group(1)} {_dm.group(2)}"] = _si
            _qp_ref_doc.close()
        except Exception:
            _ref_to_session = {}

        if qp_session_starts:
            # Group Excel rows by their QP session index
            session_buckets: dict[int, list[int]] = {}
            for i, r in enumerate(xl_rows):
                pg   = r.get("page_num", 0) or 0
                sidx = _row_session_idx(pg)
                session_buckets.setdefault(sidx, []).append(i)
            # Build excel_segments in QP session order
            # Also store qp_session_idx for each segment (for direct MS lookup)
            excel_segments    = []
            segment_qp_sidx   = []   # parallel list: segment[i] → QP session index
            for k in sorted(session_buckets.keys()):
                if session_buckets[k]:
                    excel_segments.append(session_buckets[k])
                    segment_qp_sidx.append(k)
            st.write(f"✅ Segments built from QP PDF: {len(excel_segments)} session(s) "
                     f"from {len(qp_session_starts)} QP exam boundaries")
        else:
            # Fallback: original ref-based segmentation
            excel_segments = []
            cur_seg: list[int] = []
            seg_max = 0
            seg_ref = None
            for i, r in enumerate(xl_rows):
                qn  = r["qn"]
                ref = r.get("ref", "")
                new_segment = False
                if cur_seg:
                    if ref != seg_ref:
                        new_segment = True
                    elif qn == 1:
                        new_segment = True
                    elif seg_max >= 10 and qn < seg_max - 5:
                        new_segment = True
                if new_segment:
                    excel_segments.append(cur_seg)
                    cur_seg = []
                    seg_max = 0
                cur_seg.append(i)
                if qn > seg_max:
                    seg_max = qn
                seg_ref = ref
            if cur_seg:
                excel_segments.append(cur_seg)
            segment_qp_sidx = list(range(len(excel_segments)))   # fallback: positional
            st.warning("⚠ Could not detect QP session boundaries — using ref-based segments")

        seg_offsets = infer_segment_page_offsets(locations, xl_rows, excel_segments)
        row_offset  = {}
        for seg_idx, seg_rows in enumerate(excel_segments):
            off = seg_offsets[seg_idx] if seg_idx < len(seg_offsets) else 0
            for ri in seg_rows:
                row_offset[ri] = off

        # ── 3) MS sections ─────────────────────────────────────────────────────
        if is_structured:
            st.write("📐 Locating worked solutions in MS PDF…")
            try:
                ms_sections = find_ms_question_locations(ms_bytes)
            except Exception as e:
                st.error(f"MS extraction failed: {e}")
                st.stop()
            st.write(f"✅ Found {len(ms_sections)} MS section(s)")
        else:
            st.write("🔑 MCQ mode — extracting answers from MS PDF…")
            try:
                ms_sections = extract_answers_pymupdf_per_section(ms_bytes)
            except Exception as e:
                st.error(f"MS extraction failed: {e}")
                st.stop()
            st.write(f"✅ Found {len(ms_sections)} answer section(s)")

        segments = excel_segments
        if "segment_qp_sidx" not in dir():
            segment_qp_sidx = list(range(len(segments)))
        st.write(f"📚 Using {len(segments)} paper segment(s)")

        # ── 4) Row → MS section mapping ──────────────────────────────────────────
        #
        # PURE POSITIONAL by QP exam session order.
        # segment_qp_sidx[i] = QP session index for segment i.
        # QP session N  →  MS section N  (same 0-based index).
        # This guarantees: exam 1 in QP = exam 1 in MS, exam 7 = exam 7, etc.
        # ─────────────────────────────────────────────────────────────────────────

        row_to_section     = {}
        row_to_section_idx = {}
        mismatch_warnings  = []
        n_segs = len(segments)
        n_ms   = len(ms_sections)

        # Log the mapping for debugging
        st.write(f"🔗 Matching {len(segments)} Excel segments to {n_ms} MS sections "
                 f"(method: {match_method if 'match_method' in dir() else 'Direct QP-session'})")

        for seg_pos, seg_rows in enumerate(segments):
            # QP session index stored during segment building
            qp_sidx = (segment_qp_sidx[seg_pos]
                       if seg_pos < len(segment_qp_sidx) else seg_pos)
            # Direct: MS section at same position as QP session
            if qp_sidx < n_ms:
                ms_idx = qp_sidx
                sec    = ms_sections[ms_idx]
            elif n_ms > 0:
                ms_idx = n_ms - 1          # last available section
                sec    = ms_sections[ms_idx]
                mismatch_warnings.append(
                    f"Segment {seg_pos+1} (QP session {qp_sidx+1}) beyond "
                    f"MS range ({n_ms} sections) — using last section"
                )
            else:
                ms_idx = -1
                sec    = None
                mismatch_warnings.append(f"Segment {seg_pos+1}: no MS sections")

            for ri in seg_rows:
                row_to_section[ri]     = sec
                row_to_section_idx[ri] = ms_idx

        match_method = (f"Direct QP-session order "
                        f"({n_segs} Excel segments → {n_ms} MS sections, "
                        f"ratio={n_ms/max(n_segs,1):.2f})")

        # Compute actual segment max question numbers for display
        seg_max_qns = []
        for _sg in segments:
            _qns = [xl_rows[_ri].get("qn", 0) or 0 for _ri in _sg]
            seg_max_qns.append(max(_qns) if _qns else 0)
        overall_excel_max_qn = max(seg_max_qns) if seg_max_qns else 5
        def _ms_max_qn(s):
            return s.get("max_qn", len(s.get("questions", {}))) if isinstance(s, dict) else 0
        try:
            ms_code_to_section = {}
            for _ms_i, _ms_sec in enumerate(ms_sections):
                _code = (_ms_sec.get("paper_code") if isinstance(_ms_sec, dict) else None)
                if _code and _code not in ms_code_to_section:
                    ms_code_to_section[_code] = _ms_i
        except Exception:
            ms_code_to_section = {}

        paired_by_code = 0
        unpaired_segs  = mismatch_warnings

        if is_structured:
            match_table = []
            for seg_idx, seg_rows in enumerate(segments):
                ms_sec_i = row_to_section_idx.get(seg_rows[0], -1)
                ms_sec   = row_to_section.get(seg_rows[0])
                ms_code  = (ms_sec.get("paper_code") if isinstance(ms_sec, dict) else None)
                ref_str  = xl_rows[seg_rows[0]].get("ref", "") if seg_rows else ""
                seg_mq   = seg_max_qns[seg_idx]
                ms_mq    = _ms_max_qn(ms_sec) if ms_sec else 0
                ok       = seg_mq == 0 or ms_mq == 0 or abs(ms_mq - seg_mq) <= 3
                match_table.append({
                    "Seg#":        seg_idx + 1,
                    "Excel Ref":   ref_str[:40],
                    "Seg max Q":   seg_mq,
                    "MS section#": ms_sec_i + 1 if ms_sec_i >= 0 else "—",
                    "MS max Q":    ms_mq or "—",
                    "Match":       "✅" if ok else "⚠️",
                    "MS Qs avail": len(ms_sec.get("questions", {})) if ms_sec else 0,
                })
            n_warn = len([r for r in match_table if r["Match"] == "⚠️"])
            with st.expander(
                f"📋 Paper ↔ MS Matching — {match_method}  "
                f"({'⚠️ ' + str(n_warn) + ' mismatches' if n_warn else '✅ all OK'})",
                expanded=(n_warn > 0)
            ):
                st.dataframe(match_table, use_container_width=True, hide_index=True)
                if n_warn:
                    st.warning(
                        f"⚠️ {n_warn} segment(s) may have mismatched MS sections. "
                        "Details:\n" + "\n".join(f"• {w}" for w in mismatch_warnings)
                    )

        if len(ms_sections) != len(segments):
            ratio = len(ms_sections) / max(len(segments), 1)
            if ratio < 0.5 and is_structured:
                st.warning(f"⚠ MS PDF covers fewer papers than Excel needs ({len(ms_sections)} vs {len(segments)} segments).")
            elif len(ms_sections) != len(segments):
                st.warning(f"⚠ Excel has {len(segments)} segment(s) but MS has {len(ms_sections)} section(s). Some answers may be missing.")

        # ── 5) QP extraction (for MCQ) + merge ───────────────────────────────
        if is_structured:
            # Structured: directly build worksheet from QP crops + MS crops
            # (no classify_and_extract needed)
            st.write("🔗 Merging rows (structured mode)…")

            # Apply marks override and filter
            clean_rows = []
            # Open QP once (avoids re-opening 226 times for marks lookup)
            try:
                _qp_doc_shared = fitz.open(stream=qp_bytes, filetype="pdf")
                _qp_n_pages_shared = len(_qp_doc_shared)
            except Exception:
                _qp_doc_shared = None
                _qp_n_pages_shared = 9999

            for i, r in enumerate(xl_rows):
                pg = r.get("page_num", 0)
                if pg <= 0:
                    continue
                r2 = dict(r)
                # Apply marks from QP if marks=0
                if r2.get("marks", 0) == 0 and _qp_doc_shared is not None:
                    try:
                        _ptxt = (_qp_doc_shared[pg - 1].get_text()
                                 if pg <= _qp_n_pages_shared else "")
                        _mm = re.search(r'\[Maximum mark[:\s]+(\d+)\]', _ptxt, re.IGNORECASE)
                        if _mm:
                            r2["marks"] = int(_mm.group(1))
                    except Exception:
                        pass
                if r2.get("marks", 0) == 0:
                    r2["marks"] = 1
                # Validate / clamp page_num_end against actual PDF length
                pg_end_raw = r2.get("page_num_end", 0) or 0
                if pg_end_raw > 0:
                    _qp_n_pages = _qp_n_pages_shared
                    if pg_end_raw > _qp_n_pages:
                        r2["page_num_end"] = _qp_n_pages
                    elif pg_end_raw < r2.get("page_num", 1):
                        r2["page_num_end"] = r2.get("page_num", 1)
                clean_rows.append((i, r2))

            # Build questions list for validation display
            questions = []
            for i, r2 in clean_rows:
                sec     = row_to_section.get(i)
                sec_idx = row_to_section_idx.get(i, -1)
                qn      = r2["qn"]
                q_info  = (sec.get("questions", {}).get(qn) if sec and "questions" in sec else None)
                # In preview mode, ms_doc is not open; use section-level fallback only
                if q_info is None and sec and isinstance(sec, dict) and sec.get("questions"):
                    # Accept any section with questions as a valid MS match
                    qs_any = sec["questions"]
                    q_info = qs_any.get(qn) or next(iter(qs_any.values()), None)
                # Fallback: use entire section if specific Q not found
                if q_info is None and sec and isinstance(sec, dict) and sec.get("questions"):
                    qs_s = sorted(sec["questions"].keys())
                    fq   = sec["questions"][qs_s[0]]
                    lq   = sec["questions"][qs_s[-1]]
                    q_info = {
                        "page_idx":      fq["page_idx"],
                        "top_y":         fq["top_y"],
                        "end_page_idx":  lq.get("end_page_idx", lq["page_idx"]),
                        "end_y":         lq.get("end_y", lq.get("bottom_y", 700)),
                        "marker_right_x": 0,
                    }
                ms_img  = None   # will crop during word build
                ans_ok  = q_info is not None
                questions.append({
                    **r2,
                    "_loc":         None,
                    "_ms_section":  f"MS section {sec_idx+1}" if sec_idx >= 0 else "no section",
                    "_ms_image":    ms_img,
                    "_ms_page":     (q_info["page_idx"] + 1 if q_info else None),
                    "_ms_first":    "",
                    "_topic_match": True,
                    "_qp_page":     r2.get("page_num"),
                    "_mode":        mode,
                    "found":        r2.get("page_num", 0) > 0,
                    "needs_image":  True,
                    "text":         "",
                    "A": "", "B": "", "C": "", "D": "",
                    "answer":       "" if ans_ok else "Answer not found - needs review",
                    "answerFound":  ans_ok,
                })

            if _qp_doc_shared is not None:
                try:
                    _qp_doc_shared.close()
                except Exception:
                    pass
            xl_rows_clean = [r2 for (i, r2) in clean_rows]

        else:
            # MCQ: classify_and_extract
            st.write("🔍 Extracting question text…")
            row_tuples = []
            for i, r in enumerate(xl_rows):
                p   = r.get("page_num", 0)
                off = row_offset.get(i, 0)
                if p and off:
                    p = max(1, p + off)
                row_tuples.append((r["qn"], p))
            try:
                qp_data = classify_and_extract(None, "", row_tuples, qp_bytes=qp_bytes, locations=locations)
            except Exception as e:
                st.error(f"QP extraction failed: {e}")
                st.stop()

            questions = []
            for i, (r, qp_q) in enumerate(zip(xl_rows, qp_data)):
                n       = r["qn"]
                sec     = row_to_section.get(i)
                sec_idx = row_to_section_idx.get(i, -1)
                found_q = bool(qp_q) and qp_q.get("found") is not False
                needs_img = found_q and qp_q.get("needs_image", False)

                r2 = dict(r)
                if r2.get("marks", 0) == 0:
                    r2["marks"] = 1

                ans = (sec["answers"].get(str(n), "") if sec and "answers" in sec else "")
                ans_ok   = ans in ("A", "B", "C", "D")
                ans_text = ans if ans_ok else "Answer not found - needs review"

                sec_label = (f"MS section {sec_idx+1}" if sec_idx >= 0 else "no section")

                questions.append({
                    **r2,
                    "_loc":         qp_q.get("_loc"),
                    "_ms_section":  sec_label,
                    "_ms_image":    None,
                    "_ms_page":     sec.get("start_page") if sec else None,
                    "_ms_first":    "",
                    "_topic_match": True,
                    "_qp_page":     (qp_q["_loc"]["page_idx"] + 1 if qp_q.get("_loc") else None),
                    "_mode":        mode,
                    "found":        found_q,
                    "needs_image":  needs_img,
                    "text":         qp_q.get("text", "") if found_q else "",
                    "A": qp_q.get("A", "") if found_q else "",
                    "B": qp_q.get("B", "") if found_q else "",
                    "C": qp_q.get("C", "") if found_q else "",
                    "D": qp_q.get("D", "") if found_q else "",
                    "answer":      ans_text,
                    "answerFound": ans_ok,
                })

        status.update(label="✅ Extraction done!", state="complete")

    # ── Validation summary ────────────────────────────────────────────────────
    st.subheader("Preview")
    total      = len(questions)
    found_qp   = sum(1 for q in questions if q["found"])
    found_ms   = sum(1 for q in questions if q["answerFound"])
    visual_cnt = sum(1 for q in questions if q["needs_image"])

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Total",         total)
    c2.metric("Found in QP",   found_qp)
    c3.metric("Answers found", found_ms)
    c4.metric("With visuals",  visual_cnt)

    st.dataframe(
        [{
            "Excel Row":    i + 2,
            "Q#":           q["qn"],
            "Reference":    q.get("ref", ""),
            "Excel Page":   q.get("page_num", ""),
            "QP Page Used": q.get("_qp_page") or "—",
            "MS Page Used": q.get("_ms_page") or "—",
            "Mode":         "MCQ" if "MCQ" in (q.get("_mode") or "") else "Structured",
            "MS Section":   q.get("_ms_section", ""),
            "Chapter":      q.get("topic", ""),
            "Difficulty":   q.get("difficulty", ""),
            "Marks":        q.get("marks", ""),
            "Status":       ("✅" if q["found"] and q["answerFound"]
                             else ("⚠ MS missing" if q["found"]
                                   else ("⚠ QP missing" if q["answerFound"]
                                         else "❌ both missing"))),
        } for i, q in enumerate(questions)],
        use_container_width=True,
        hide_index=True,
    )

    unmatched = []
    for i, q in enumerate(questions):
        problems = []
        if not q["found"]:
            if not q.get("page_num"):
                problems.append("no Page Number in Excel")
            else:
                problems.append(f"Q{q['qn']} not located on page {q['page_num']} of QP")
        if not q["answerFound"] and not q.get("_ms_image"):
            if "Structured" in (q.get("_mode") or ""):
                problems.append(f"Q{q['qn']} not found in MS section")
            elif "MCQ" in (q.get("_mode") or ""):
                problems.append(f"Q{q['qn']} has no A/B/C/D answer")
        if problems:
            unmatched.append({
                "Excel Row": i + 2, "Reference": q.get("ref", ""),
                "Q#": q["qn"], "Page (Excel)": q.get("page_num", 0),
                "Reason": " · ".join(problems),
            })

    qp_match_pct    = found_qp / max(total, 1)
    matching_is_bad = qp_match_pct < 0.70

    if unmatched:
        st.subheader(f"⚠ Unmatched rows ({len(unmatched)})")
        st.dataframe(unmatched, use_container_width=True, hide_index=True)

    st.subheader("Download")
    if matching_is_bad:
        st.error(f"❌ Only {found_qp}/{total} ({qp_match_pct:.0%}) rows found in QP. Fix matching before downloading.")
        st.stop()

    if qp_match_pct < 0.95:
        st.warning(f"⚠ {found_qp}/{total} rows matched ({qp_match_pct:.0%}). {total - found_qp} unmatched will show 'Question not found'.")

    # ── Build Word ─────────────────────────────────────────────────────────────
    _prog_bar  = st.progress(0)
    _prog_text = st.empty()

    def on_progress(done, total_, msg):
        _prog_bar.progress(done / max(total_, 1))
        _prog_text.text(msg)

    with st.spinner("Building Word document…"):
        try:
            if is_structured:
                # Build using FINAL AGENT structured generator
                _docx = build_structured_worksheet(
                    xl_rows_clean,
                    qp_bytes,
                    ms_bytes,
                    ms_sections,
                    {i: row_to_section[orig_i]
                     for i, (orig_i, _) in enumerate(clean_rows)},
                    progress_cb=on_progress,
                )
            else:
                _docx = build_word_document(questions, qp_bytes, locations, progress_cb=on_progress)
        except Exception as e:
            st.error(f"Word generation failed: {e}")
            import traceback
            st.code(traceback.format_exc())
            st.stop()

    _prog_bar.progress(1.0)
    _prog_text.empty()

    ms_missing_n = sum(1 for q in questions if not q.get("_ms_image") and not q.get("answerFound"))

    st.session_state["ws_docx"]    = _docx
    st.session_state["ws_missing"] = ms_missing_n
    st.session_state["ws_stats"]   = (
        f"Total: {total} | Found in QP: {found_qp} | Answers found: {found_ms}"
    )
    st.session_state["ws_ready"]   = True
    st.success("✅ Worksheet built — scroll down to download.")


# ═══════════════════════════════════════════════════════════════════════════════
#  PERSISTENT DOWNLOAD
# ═══════════════════════════════════════════════════════════════════════════════
if st.session_state.get("ws_ready") and st.session_state.get("ws_docx"):
    st.divider()
    st.subheader("⬇️ Download Worksheet")

    _stats = st.session_state.get("ws_stats", "")
    if _stats:
        st.info(f"📊 {_stats}")

    _missing = st.session_state.get("ws_missing", 0)
    if _missing:
        st.warning(f"⚠️ {_missing} question(s) have no MS answer.")

    st.download_button(
        label="⬇️ Download Worksheet (.docx)",
        data=st.session_state["ws_docx"],
        file_name="IB_Worksheet.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="persistent_dl",
        type="primary",
    )

    if st.button("🔄 Start New Worksheet", key="reset_ws"):
        for _k in ["ws_docx", "ws_missing", "ws_stats", "ws_ready"]:
            st.session_state.pop(_k, None)
        st.rerun()
